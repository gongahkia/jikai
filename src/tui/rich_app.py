"""Jikai Rich TUI - Procedural terminal interface (caipng-style)."""

import asyncio
import csv
from dataclasses import dataclass
import json
import logging
import os
from pathlib import Path
import re
from typing import Dict, List, Optional

import questionary
from questionary import Choice
from rich import box
from rich.console import Console
from rich.live import Live
from rich.panel import Panel
from rich.progress import (
    BarColumn,
    Progress,
    SpinnerColumn,
    TextColumn,
    TimeElapsedColumn,
)
from rich.table import Table
from rich.text import Text

console = Console()
LOG_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "logs")
os.makedirs(LOG_DIR, exist_ok=True)
LOG_PATH = os.path.join(LOG_DIR, "tui.log")
_fh = logging.FileHandler(LOG_PATH, mode="w", encoding="utf-8")
_fh.setFormatter(
    logging.Formatter("%(asctime)s  %(message)s", datefmt="%Y-%m-%d %H:%M:%S")
)
tlog = logging.getLogger("jikai.tui")
tlog.setLevel(logging.INFO)
tlog.addHandler(_fh)

TOPICS = [
    "negligence",
    "duty_of_care",
    "causation",
    "remoteness",
    "battery",
    "assault",
    "false_imprisonment",
    "defamation",
    "private_nuisance",
    "trespass_to_land",
    "vicarious_liability",
    "strict_liability",
    "harassment",
    "occupiers_liability",
    "product_liability",
    "contributory_negligence",
    "economic_loss",
    "psychiatric_harm",
    "employers_liability",
]

TOPIC_DESCRIPTIONS = {
    "negligence": "duty, breach, damage, causation",
    "duty_of_care": "neighbour principle, proximity",
    "causation": "but-for test, legal causation",
    "remoteness": "foreseeability of damage",
    "contributory_negligence": "claimant's own fault",
    "battery": "intentional application of force",
    "assault": "apprehension of immediate contact",
    "false_imprisonment": "unlawful restraint of liberty",
    "trespass_to_land": "unlawful entry onto land",
    "vicarious_liability": "employer liability for employee torts",
    "strict_liability": "liability without fault",
    "occupiers_liability": "duties to visitors and trespassers",
    "employers_liability": "workplace safety duties",
    "product_liability": "defective product claims",
    "defamation": "false statements harming reputation",
    "private_nuisance": "unreasonable interference with land use",
    "harassment": "course of conduct causing alarm",
    "economic_loss": "pure financial loss claims",
    "psychiatric_harm": "nervous shock and mental injury",
}

TOPIC_CATEGORIES = {
    "Negligence-Based": ["negligence", "duty_of_care", "causation", "remoteness", "contributory_negligence"],
    "Intentional Torts": ["battery", "assault", "false_imprisonment", "trespass_to_land"],
    "Liability": ["vicarious_liability", "strict_liability", "occupiers_liability", "employers_liability", "product_liability"],
    "Specific Torts": ["defamation", "private_nuisance", "harassment"],
    "Damages": ["economic_loss", "psychiatric_harm"],
}

PROVIDERS = ["ollama", "openai", "anthropic", "google", "local"]


def _topic_choices():
    """Build topic choices with descriptions, grouped by category."""
    choices = []
    for category, topics in TOPIC_CATEGORIES.items():
        # Add category separator
        choices.append(Choice(f"── {category} ──", value=None, disabled=""))
        for t in topics:
            desc = TOPIC_DESCRIPTIONS.get(t, "")
            title = t.replace("_", " ").title()
            label = f"{title} — {desc}" if desc else title
            choices.append(Choice(label, value=t))
    return choices


COMPLEXITY_CHOICES = [
    Choice("1 — Simple", value="1"),
    Choice("2 — Basic", value="2"),
    Choice("3 — Moderate", value="3"),
    Choice("4 — Complex", value="4"),
    Choice("5 — Advanced", value="5"),
]
PARTIES_CHOICES = [
    Choice("2 parties", value="2"),
    Choice("3 parties", value="3"),
    Choice("4 parties", value="4"),
    Choice("5 parties", value="5"),
]
METHOD_CHOICES = [
    Choice("Pure LLM", value="pure_llm"),
    Choice("ML Assisted", value="ml_assisted"),
    Choice("Hybrid", value="hybrid"),
]
PROVIDER_CHOICES = [
    Choice("Ollama (local)", value="ollama"),
    Choice("OpenAI", value="openai"),
    Choice("Anthropic", value="anthropic"),
    Choice("Google", value="google"),
    Choice("Local LLM", value="local"),
]
PROVIDER_MODELS = {
    "ollama": ["llama3", "llama3.1", "mistral", "gemma2", "phi3", "codellama", "qwen2"],
    "openai": [
        "gpt-4o",
        "gpt-4o-mini",
        "gpt-4-turbo",
        "gpt-3.5-turbo",
        "o1",
        "o1-mini",
    ],
    "anthropic": [
        "claude-sonnet-4-5-20250929",
        "claude-opus-4-20250918",
        "claude-haiku-4-5-20251001",
        "claude-3-5-sonnet-20241022",
    ],
    "google": ["gemini-2.0-flash", "gemini-1.5-pro", "gemini-1.5-flash"],
    "local": ["default"],
}


def _model_choices(provider):
    choices = [Choice(m, value=m) for m in PROVIDER_MODELS.get(provider, [])]
    choices.append(Choice("Custom...", value="__custom__"))
    return choices


def _select(message, choices):
    """Arrow-key select menu. Returns value or None on Ctrl+C."""
    try:
        return questionary.select(message, choices=choices).ask()
    except KeyboardInterrupt:
        return None


_HOTKEY_MAP = {
    "g": "gen",
    "h": "history",
    "s": "settings",
    "p": "providers",
}

# Global hotkeys that work from any submenu
_GLOBAL_HOTKEYS = {
    "g": "__jump_gen__",
}

HELP_DESCRIPTIONS = {
    "ocr": "extract text from PDF/DOCX/images into the corpus (requires: pymupdf, python-docx, pillow, pytesseract)",
    "corpus": "view, search, and filter preprocessed corpus entries",
    "label": "tag corpus entries with topics and quality scores for training",
    "train": "train ML models on labelled data for better generation (requires: scikit-learn, pandas)",
    "embed": "create vector embeddings for semantic search during generation (requires: chromadb, sentence-transformers, torch)",
    "gen": "create a new tort law scenario using AI",
    "export": "save a generated hypothetical as DOCX or PDF (requires: python-docx)",
    "history": "browse past generations with search and filter",
    "stats": "view statistics about your corpus and generations",
    "tools": "batch operations: bulk generate, import cases, bulk label",
    "settings": "configure API keys, hosts, and defaults",
    "providers": "check health and set default LLM provider",
    "guided": "step-by-step walkthrough for first-time users",
    "cleanup": "selectively remove generated files, models, logs, etc.",
    "more": "access secondary features: history, stats, settings, etc.",
}

SERVICE_DEPS = {
    "ocr": ["pymupdf", "python-docx", "pillow", "pytesseract"],
    "train": ["scikit-learn", "pandas", "joblib"],
    "embed": ["chromadb", "sentence-transformers", "torch"],
    "export": ["python-docx"],
    "gen": ["httpx"],
    "google": ["google-generativeai"],
    "anthropic": ["anthropic"],
}

# Map pip package names to import names for verification
IMPORT_MAP = {
    "pymupdf": "fitz",
    "python-docx": "docx",
    "pillow": "PIL",
    "scikit-learn": "sklearn",
    "sentence-transformers": "sentence_transformers",
    "google-generativeai": "google.generativeai",
}


def _show_help_overlay(choices):
    """Show a help panel describing each visible option with contextual descriptions."""
    lines = []
    for c in choices:
        if hasattr(c, "title") and hasattr(c, "value"):
            title = c.title
            if isinstance(title, list):  # prompt_toolkit formatted text tuples
                title = "".join(t[1] for t in title)
            value = c.value
            desc = HELP_DESCRIPTIONS.get(value, "")
            disabled = getattr(c, "disabled", None)
            status = f" [dim]({disabled})[/dim]" if disabled else ""
            if desc:
                lines.append(f"  [cyan]{title}[/cyan]{status}\n    [dim]{desc}[/dim]")
            else:
                lines.append(f"  [cyan]{title}[/cyan]{status}")
    help_text = "\n".join(lines) if lines else "No options available."
    # Add getting started section for first-time users
    getting_started = (
        "[bold]Getting Started:[/bold]\n"
        "  [dim]New here? Start with Import & Preprocess Corpus, then Generate Hypothetical.[/dim]"
    )
    console.print(
        Panel(
            f"[bold]Available Options:[/bold]\n{help_text}\n\n"
            f"{getting_started}\n\n"
            "[dim]Shortcuts: q=quit, g=generate, h=history, s=settings, p=providers[/dim]\n"
            "[dim]Press ? for this help. Press any key to dismiss.[/dim]",
            title="Help",
            box=box.ROUNDED,
            border_style="yellow",
        )
    )


def _select_quit(message, choices, hotkeys=None, style=None, enable_global_hotkeys=True):
    """Arrow-key select with q-to-quit, ?-for-help, and optional hotkey bindings."""
    from prompt_toolkit.key_binding import KeyBindings, merge_key_bindings

    hk = hotkeys or _HOTKEY_MAP
    _result = {"value": None}
    kb = KeyBindings()

    @kb.add("q")
    def _quit_on_q(event):
        event.app.exit(exception=KeyboardInterrupt())

    @kb.add("?")
    def _help_on_question(event):
        _result["value"] = "__help__"
        event.app.exit(result="__help__")

    # Build hotkeys that return corresponding menu values
    choice_values = {c.value for c in choices if hasattr(c, "value")}
    for key, val in hk.items():
        if val in choice_values:

            def _make_handler(v):
                def _handler(event):
                    _result["value"] = v
                    event.app.exit(result=v)

                return _handler

            kb.add(key)(_make_handler(val))

    # Add global hotkeys that work from any submenu (e.g., 'g' to jump to Generate)
    if enable_global_hotkeys:
        for key, val in _GLOBAL_HOTKEYS.items():
            if key not in [k for k, v in hk.items() if v in choice_values]:
                def _make_global_handler(v):
                    def _handler(event):
                        _result["value"] = v
                        event.app.exit(result=v)
                    return _handler
                kb.add(key)(_make_global_handler(val))

    hint_keys = " ".join(f"{k}={v}" for k, v in hk.items() if v in choice_values)
    global_hint = "g=generate" if enable_global_hotkeys and "gen" not in choice_values else ""
    all_hints = " | ".join(filter(None, [hint_keys, global_hint]))
    instruction = (
        f"(q quit | ? help | {all_hints})" if all_hints else "(q quit | ? help)"
    )
    while True:
        try:
            _result["value"] = None
            q = questionary.select(
                message, choices=choices, instruction=instruction, style=style
            )
            orig = q.application.key_bindings
            q.application.key_bindings = merge_key_bindings([orig, kb]) if orig else kb
            result = q.unsafe_ask()
            chosen = _result["value"] if _result["value"] is not None else result
            if chosen == "__help__":
                _show_help_overlay(choices)
                continue
            return chosen
        except (KeyboardInterrupt, EOFError):
            return None


def _checkbox(message, choices):
    """Arrow-key checkbox menu. Returns list or None on Ctrl+C."""
    try:
        return questionary.checkbox(message, choices=choices).ask()
    except KeyboardInterrupt:
        return None


def _text(message, default=""):
    """questionary text input with default."""
    try:
        return questionary.text(message, default=default).ask()
    except KeyboardInterrupt:
        return None


def _password(message, default=""):
    """questionary password input (masked)."""
    try:
        return questionary.password(message, default=default).ask()
    except KeyboardInterrupt:
        return None


def _confirm(message, default=True):
    """questionary yes/no confirm."""
    try:
        return questionary.confirm(message, default=default).ask()
    except KeyboardInterrupt:
        return False


def _path(message, default=""):
    """questionary path input with autocomplete."""
    try:
        return questionary.path(message, default=default).ask()
    except KeyboardInterrupt:
        return None


def _validate_number(min_val, max_val, is_float=False):
    """Return a questionary validator for numeric range."""

    def _validator(val):
        try:
            n = float(val) if is_float else int(val)
        except (ValueError, TypeError):
            kind = "float" if is_float else "integer"
            return f"Must be a valid {kind}"
        if n < min_val or n > max_val:
            return f"Must be between {min_val} and {max_val}"
        return True

    return _validator


def _validated_text(message, default="", validate=None):
    """questionary text with validation."""
    try:
        return questionary.text(message, default=default, validate=validate).ask()
    except KeyboardInterrupt:
        return None


def _run_async(coro):
    """Run async coroutine from synchronous context."""
    return asyncio.run(coro)


HISTORY_PATH = "data/history.json"


@dataclass
class GenerationConfig:
    """Encapsulates all generation parameters."""

    topic: str
    provider: str
    model: Optional[str]
    complexity: int
    parties: int
    method: str
    temperature: float = 0.7
    red_herrings: bool = False


class JikaiTUI:
    def __init__(self):
        self._loaded_path = ""
        self._entries: list = []
        self._corpus_path = "corpus/clean/tort/corpus.json"
        self._raw_dir = "corpus/raw"
        self._models_dir = "models"
        self._train_data = "corpus/labelled/sample.csv"
        self._nav_path: list = []
        self._last_history_load_time: float = 0.0
        self._cached_last_score: str = "—"

    def _push_nav(self, label: str):
        self._nav_path.append(label)

    def _pop_nav(self):
        if self._nav_path:
            self._nav_path.pop()

    def _render_breadcrumb(self):
        if self._nav_path:
            crumbs = " > ".join(self._nav_path)
            console.print(f"[dim]{crumbs}[/dim]")

    def display_banner(self):
        tlog.info("=== TUI session started ===")
        console.print(
            Panel(
                "[bold cyan]Jikai - Legal Hypothetical Generator[/bold cyan]\n"
                "[dim]AI-powered Singapore tort law hypothetical generation[/dim]",
                box=box.DOUBLE,
                border_style="cyan",
            )
        )

    def _corpus_ready(self):
        return Path(self._corpus_path).exists()

    def _labelled_ready(self):
        return Path(self._train_data).exists()

    def _models_ready(self):
        md = Path(self._models_dir)
        return all(
            (md / f).exists()
            for f in ("classifier.joblib", "vectorizer.joblib", "binarizer.joblib")
        )

    def _embeddings_ready(self):
        return Path("chroma_db").exists() and any(Path("chroma_db").iterdir())

    def _render_status_bar(self):
        """Render a persistent contextual status bar."""
        state = self._load_state()
        last_cfg = state.get("last_config", {})
        provider = last_cfg.get("provider", "—")
        model = last_cfg.get("model", "—")
        corpus_count = 0
        if self._corpus_ready():
            try:
                with open(self._corpus_path) as f:
                    data = json.load(f)
                corpus_count = (
                    len(data)
                    if isinstance(data, list)
                    else len(data.get("entries", []))
                )
            except Exception:
                pass
        # Build status parts - only show what's relevant
        parts = [
            f"[dim]Provider:[/dim] {provider}/{model or '—'}",
            f"[dim]Corpus:[/dim] {corpus_count}",
        ]
        # Only show ML Models if user has trained any
        models_ready = self._models_ready()
        if models_ready or self._labelled_ready():
            icon = "[green]✓ trained[/green]" if models_ready else "[dim]○ not trained[/dim]"
            parts.append(f"[dim]ML Models:[/dim] {icon}")
        # Only show Semantic Search if user has embedded or has corpus
        embed_ready = self._embeddings_ready()
        if embed_ready or corpus_count > 0:
            icon = "[green]✓ indexed[/green]" if embed_ready else "[dim]○ off[/dim]"
            parts.append(f"[dim]Semantic Search:[/dim] {icon}")
        # Last quality score — cached for 5s
        import time as _time

        now = _time.time()
        if now - self._last_history_load_time > 5.0:
            history = self._load_history()
            self._cached_last_score = "—"
            if history:
                s = history[-1].get("validation_score", history[-1].get("score"))
                if s is not None:
                    self._cached_last_score = f"{float(s):.1f}"
            self._last_history_load_time = now
        last_score = self._cached_last_score
        if last_score != "—":
            parts.append(f"[dim]Last Score:[/dim] {last_score}")
        console.print(
            Panel(
                "  ".join(parts),
                box=box.SIMPLE,
                border_style="dim",
            )
        )

    def main_menu(self):
        self._nav_path = ["Main"]
        while True:
            self._render_breadcrumb()
            self._render_status_bar()
            console.print("\n[bold yellow]Main Menu[/bold yellow]")
            corpus_ok = self._corpus_ready()
            labelled_ok = self._labelled_ready()
            models_ok = self._models_ready()
            embed_ok = self._embeddings_ready()

            # Dep checks
            ocr_deps = self._check_service_deps("ocr")
            train_deps = self._check_service_deps("train")
            embed_deps = self._check_service_deps("embed")
            gen_deps = self._check_service_deps("gen")
            export_deps = self._check_service_deps("export")

            browse_disabled = None if corpus_ok else "preprocess corpus first"
            label_disabled = None if corpus_ok else "preprocess corpus first"

            # Complex disabled logic with dep check
            if not corpus_ok:
                train_disabled = "preprocess corpus first"
            elif not labelled_ok:
                train_disabled = "label corpus first"
            elif not train_deps:
                train_disabled = "install ML dependencies"
            else:
                train_disabled = None

            if not corpus_ok:
                embed_disabled = "preprocess corpus first"
            elif not embed_deps:
                embed_disabled = "install embedding dependencies"
            else:
                embed_disabled = None

            if not corpus_ok:
                gen_disabled = "preprocess corpus first"
            elif not gen_deps:
                gen_disabled = "install generation dependencies"
            else:
                gen_disabled = None

            history = self._load_history()
            has_history = len(history) > 0
            if not has_history:
                export_disabled = "generate a hypothetical first"
            elif not export_deps:
                export_disabled = "install export dependencies"
            else:
                export_disabled = None

            from prompt_toolkit.styles import Style as PtStyle

            menu_style = PtStyle.from_dict({
                "ok": "#00aa00",
                "fail": "#cc0000",
                "dim": "#666666",
                "warn": "#ffff00",
            })

            def _tag(ok, optional=False, deps_ok=True):
                tags = []
                if not deps_ok:
                    tags.append(("class:warn", " ⚠"))
                if ok:
                    tags.append(("class:ok", " ✓"))
                else:
                    tags.append(("class:dim", " ○") if optional else ("class:fail", " ✗"))
                return tags

            def _lbl(text, enabled=True):
                cls = "" if enabled else "class:dim"
                return [(cls, text)]

            choice = _select_quit(
                "What would you like to do?",
                choices=[
                    Choice(
                        title=_lbl("Import & Preprocess Corpus") + _tag(corpus_ok, deps_ok=ocr_deps),
                        value="ocr",
                    ),
                    Choice(
                        title=_lbl("Browse Corpus", corpus_ok),
                        value="corpus",
                        disabled=browse_disabled,
                    ),
                    Choice(
                        title=_lbl("Label Corpus") + _tag(labelled_ok),
                        value="label",
                        disabled=label_disabled,
                    ),
                    Choice(
                        title=_lbl("Train ML Models (optional)", labelled_ok and train_deps)
                        + _tag(models_ok, optional=True, deps_ok=train_deps),
                        value="train",
                        disabled=train_disabled,
                    ),
                    Choice(
                        title=_lbl("Index Semantic Search (optional)", corpus_ok and embed_deps)
                        + _tag(embed_ok, optional=True, deps_ok=embed_deps),
                        value="embed",
                        disabled=embed_disabled,
                    ),
                    Choice(
                        title=_lbl("Generate Hypothetical", corpus_ok and gen_deps)
                        + _tag(corpus_ok, deps_ok=gen_deps),
                        value="gen",
                        disabled=gen_disabled,
                    ),
                    Choice(
                        title=_lbl("Export to DOCX / PDF", has_history and export_deps)
                        + _tag(has_history, deps_ok=export_deps),
                        value="export",
                        disabled=export_disabled,
                    ),
                    Choice("More... ›", value="more"),
                ],
                style=menu_style,
            )
            if choice is None:
                if _confirm("Clean up Jikai files before exiting?", default=False):
                    self._push_nav("Clean Up")
                    self.cleanup_flow()
                    self._pop_nav()
                console.print("[dim]Goodbye.[/dim]")
                tlog.info("=== TUI session ended ===")
                break
            _flow_labels = {
                "ocr": "Import & Preprocess",
                "corpus": "Browse Corpus",
                "label": "Label Corpus",
                "train": "Train ML Models",
                "embed": "Index Semantic Search",
                "gen": "Generate",
                "export": "Export",
                "more": "More",
                "tools": "Batch Operations",
                "history": "History",
                "stats": "Stats",
                "settings": "Settings",
                "providers": "Providers",
            }
            if choice in _flow_labels:
                self._push_nav(_flow_labels[choice])
            if choice == "ocr":
                self.ocr_flow()
            elif choice == "corpus":
                self.corpus_flow()
            elif choice == "label":
                self.label_flow()
            elif choice == "train":
                self.train_flow()
            elif choice == "embed":
                self.embed_flow()
            elif choice == "gen":
                self.generate_flow()
            elif choice == "export":
                self.export_flow()
            elif choice == "more":
                result = self._more_menu()
                if result == "__jump_gen__":
                    if choice in _flow_labels:
                        self._pop_nav()
                    self._push_nav("Generate")
                    self.generate_flow()
                    self._pop_nav()
                    continue
            elif choice == "tools":
                self._tools_menu()
            elif choice == "history":
                self.history_flow()
            elif choice == "stats":
                self.stats_flow()
            elif choice == "settings":
                self.settings_flow()
            elif choice == "providers":
                self.providers_flow()
            # Pop nav after returning from submenu
            if choice in _flow_labels:
                self._pop_nav()

    def _more_menu(self):
        """Submenu for secondary features: History, Stats, Batch Ops, Settings, Providers, Clean Up."""
        while True:
            console.print("\n[bold yellow]More[/bold yellow]")
            c = _select_quit(
                "More options",
                choices=[
                    Choice("History", value="history"),
                    Choice("Stats Dashboard", value="stats"),
                    Choice("Batch Operations ›", value="tools"),
                    Choice("Settings", value="settings"),
                    Choice("Providers", value="providers"),
                    Choice("Guided Mode", value="guided"),
                    Choice("Clean Up", value="cleanup"),
                ],
            )
            if c is None:
                return
            if c == "__jump_gen__":
                return "__jump_gen__"  # Propagate up to main_menu
            _labels = {
                "history": "History",
                "stats": "Stats",
                "tools": "Batch Operations",
                "settings": "Settings",
                "providers": "Providers",
                "guided": "Guided Mode",
                "cleanup": "Clean Up",
            }
            if c in _labels:
                self._push_nav(_labels[c])
            if c == "history":
                self.history_flow()
            elif c == "stats":
                self.stats_flow()
            elif c == "tools":
                result = self._tools_menu()
                if result == "__jump_gen__":
                    if c in _labels:
                        self._pop_nav()
                    return "__jump_gen__"
            elif c == "settings":
                self.settings_flow()
            elif c == "providers":
                self.providers_flow()
            elif c == "guided":
                self.guided_mode()
            elif c == "cleanup":
                self.cleanup_flow()
            if c in _labels:
                self._pop_nav()

    def guided_mode(self):
        """Guided walkthrough for first-time users: Preprocess → Generate → Export."""
        console.print("\n[bold yellow]Guided Mode — 3-Step Walkthrough[/bold yellow]")
        console.print(
            Panel(
                "[bold]Welcome to Jikai![/bold]\n\n"
                "This guided mode will walk you through the essential steps:\n"
                "  [cyan]1.[/cyan] Import & preprocess your legal corpus\n"
                "  [cyan]2.[/cyan] Generate a tort law hypothetical\n"
                "  [cyan]3.[/cyan] Export to DOCX/PDF\n\n"
                "[dim]Press any key to continue, or q to exit.[/dim]",
                title="Getting Started",
                box=box.ROUNDED,
                border_style="green",
            )
        )
        # Step 1: Check if corpus exists, otherwise guide through preprocessing
        if not self._corpus_ready():
            console.print("\n[bold cyan]Step 1/3: Import & Preprocess Corpus[/bold cyan]")
            console.print(
                "[dim]You need a corpus before generating. "
                "Let's preprocess some documents.[/dim]\n"
            )
            if _confirm("Start preprocessing now?", default=True):
                self._push_nav("Preprocess")
                self.ocr_flow()
                self._pop_nav()
            if not self._corpus_ready():
                console.print("[yellow]⚠ Corpus not ready. Exiting guided mode.[/yellow]")
                return
        else:
            console.print("\n[green]✓ Step 1/3: Corpus ready[/green]")

        # Step 2: Generate a hypothetical
        console.print("\n[bold cyan]Step 2/3: Generate Hypothetical[/bold cyan]")
        console.print(
            "[dim]Now let's generate a tort law scenario using AI.[/dim]\n"
        )
        if _confirm("Generate a hypothetical now?", default=True):
            self._push_nav("Generate")
            self.generate_flow()
            self._pop_nav()

        # Step 3: Export
        history = self._load_history()
        if history:
            console.print("\n[bold cyan]Step 3/3: Export to Document[/bold cyan]")
            console.print(
                "[dim]Export your generated hypothetical to DOCX or PDF.[/dim]\n"
            )
            if _confirm("Export now?", default=True):
                self._push_nav("Export")
                self.export_flow()
                self._pop_nav()

        console.print(
            Panel(
                "[bold green]Guided mode complete![/bold green]\n\n"
                "You've learned the core workflow. Next steps:\n"
                "  • [cyan]Label Corpus[/cyan] — improve generation quality with training data\n"
                "  • [cyan]Train ML Models[/cyan] — enhance AI with custom models\n"
                "  • [cyan]History[/cyan] — review and reuse past generations\n",
                title="All Done!",
                box=box.ROUNDED,
                border_style="green",
            )
        )

    def _tools_menu(self):
        """Submenu for power tools outside the core workflow."""
        while True:
            console.print("\n[bold yellow]Batch Operations[/bold yellow]")

            gen_deps = self._check_service_deps("gen")
            ocr_deps = self._check_service_deps("ocr")

            c = _select_quit(
                "Batch Operations",
                choices=[
                    Choice(
                        "Batch Generate" + ("" if gen_deps else " ⚠"),
                        value="batch_gen",
                        disabled=None if gen_deps else "install generation dependencies",
                    ),
                    Choice(
                        "Import SG Cases" + ("" if ocr_deps else " ⚠"),
                        value="import_cases",
                        disabled=None if ocr_deps else "install OCR dependencies",
                    ),
                    Choice("Bulk Label", value="bulk_label"),
                ],
            )
            if c is None:
                return
            if c == "__jump_gen__":
                return "__jump_gen__"  # Propagate up
            _labels = {
                "batch_gen": "Batch",
                "import_cases": "Import",
                "bulk_label": "Bulk Label",
            }
            if c in _labels:
                self._push_nav(_labels[c])
            if c == "batch_gen":
                self.batch_generate_flow()
            elif c == "import_cases":
                self.import_cases_flow()
            elif c == "bulk_label":
                self.bulk_label_flow()
            if c in _labels:
                self._pop_nav()

    # ── ocr / preprocess ──────────────────────────────────────
    def ocr_flow(self):
        console.print("\n[bold yellow]OCR / Preprocess Corpus[/bold yellow]")
        while True:
            c = _select_quit(
                "OCR Menu",
                choices=[
                    Choice("Preprocess raw corpus (TXT/PDF/PNG/DOCX)", value="1"),
                    Choice("Convert single file to TXT (OCR)", value="2"),
                ],
            )
            if c is None:
                return
            if c == "__jump_gen__":
                return "__jump_gen__"
            if c == "1":
                self._preprocess_raw()
            elif c == "2":
                self._ocr_single()

    def _ocr_single(self):
        path = _path("File to OCR (PDF/PNG/JPG/DOCX)", default="")
        if not path:
            return
        try:
            from ..services.corpus_preprocessor import extract_text, normalize_text

            with console.status("[bold green]Extracting text...", spinner="dots"):
                text = normalize_text(extract_text(Path(path)))
            if not text:
                console.print("[red]✗ No text extracted (check file format/deps)[/red]")
                return
            console.print(
                Panel(
                    text[:2000] + ("..." if len(text) > 2000 else ""),
                    title=f"Extracted from {Path(path).name}",
                    box=box.ROUNDED,
                    border_style="green",
                )
            )
            out = _path("Save TXT to (enter to skip)", default="")
            if out:
                Path(out).write_text(text, encoding="utf-8")
                console.print(f"[green]✓ Saved to {out}[/green]")
            tlog.info("OCR  %s → %d chars", path, len(text))
        except Exception as e:
            console.print(f"[red]✗ OCR failed: {e}[/red]")
            console.print(
                "[dim]Tip: ensure pytesseract and tesseract are installed[/dim]"
            )

    # ── label ──────────────────────────────────────────────────
    def label_flow(self):
        console.print("\n[bold yellow]Label Corpus → Training CSV[/bold yellow]")
        # Auto-use default corpus if it exists and has entries
        corpus_path = self._corpus_path
        if Path(corpus_path).exists():
            try:
                with open(corpus_path) as f:
                    data = json.load(f)
                entries = data if isinstance(data, list) else data.get("entries", [])
                if entries:
                    console.print(f"[dim]Using corpus: {corpus_path} ({len(entries)} entries)[/dim]")
                else:
                    corpus_path = _path("Corpus JSON to label", default=self._corpus_path)
                    if corpus_path is None:
                        return
            except Exception:
                corpus_path = _path("Corpus JSON to label", default=self._corpus_path)
                if corpus_path is None:
                    return
        else:
            corpus_path = _path("Corpus JSON to label", default=self._corpus_path)
            if corpus_path is None:
                return
        out_path = _path("Output labelled CSV", default=self._train_data)
        if out_path is None:
            return
        try:
            with open(corpus_path) as f:
                data = json.load(f)
            entries = data if isinstance(data, list) else data.get("entries", [])
        except Exception as e:
            console.print(f"[red]✗ Failed to load corpus: {e}[/red]")
            return
        if not entries:
            console.print("[red]✗ No entries in corpus[/red]")
            return
        existing = []
        skip_texts = set()
        if Path(out_path).exists():
            with open(out_path) as f:
                reader = csv.DictReader(f)
                for row in reader:
                    existing.append(row)
                    skip_texts.add(row.get("text", "")[:100])
            console.print(
                f"[dim]{len(existing)} already labelled, will append new[/dim]"
            )
        unlabelled = [e for e in entries if e.get("text", "")[:100] not in skip_texts]
        console.print(
            f"[cyan]{len(unlabelled)} entries to label ({len(entries)} total)[/cyan]"
        )
        if not unlabelled:
            console.print("[green]All entries already labelled[/green]")
            return
        difficulty_choices = [
            Choice("Easy", value="easy"),
            Choice("Medium", value="medium"),
            Choice("Hard", value="hard"),
        ]
        labelled = list(existing)
        count = 0
        for i, entry in enumerate(unlabelled):
            text = entry.get("text", "")
            topics_hint = entry.get("topic", entry.get("topics", ""))
            if isinstance(topics_hint, list):
                topics_hint = ", ".join(topics_hint)
            subtitle = f"Topics: {topics_hint}" if topics_hint else None
            panel = Panel(
                text,
                title=f"Entry {i+1}/{len(unlabelled)}",
                subtitle=subtitle,
                box=box.ROUNDED,
                border_style="cyan",
            )
            if len(text) > 1000:
                with console.pager(styles=True):
                    console.print(panel)
            else:
                console.print(panel)
            action = _select_quit(
                "Action",
                choices=[
                    Choice("Label this entry", value="label"),
                    Choice("Skip", value="skip"),
                    Choice("Save & quit", value="done"),
                ],
            )
            if action is None or action == "done":
                break
            if action == "skip":
                continue
            topics = _checkbox("Topics", choices=_topic_choices())
            if topics is None:
                break
            quality = _validated_text(
                "Quality score (0-10)",
                default="7.0",
                validate=_validate_number(0.0, 10.0, is_float=True),
            )
            if quality is None:
                break
            difficulty = _select_quit("Difficulty", choices=difficulty_choices)
            if difficulty is None:
                break
            labelled.append(
                {
                    "text": text,
                    "topic_labels": "|".join(topics),
                    "quality_score": quality,
                    "difficulty_level": difficulty,
                }
            )
            count += 1
            # Show confirmation with the label just assigned
            console.print(
                f"[green]✓ Labelled ({count} this session)[/green] "
                f"[dim]Topics: {', '.join(topics[:3])}{'...' if len(topics) > 3 else ''}, "
                f"Quality: {quality}, Difficulty: {difficulty}[/dim]"
            )
            # Offer to re-label
            if _confirm("Re-label this entry?", default=False):
                # Remove the last entry and redo
                labelled.pop()
                count -= 1
                # Re-prompt for labels
                topics = _checkbox("Topics", choices=_topic_choices())
                if topics is None:
                    break
                quality = _validated_text(
                    "Quality score (0-10)",
                    default="7.0",
                    validate=_validate_number(0.0, 10.0, is_float=True),
                )
                if quality is None:
                    break
                difficulty = _select_quit("Difficulty", choices=difficulty_choices)
                if difficulty is None:
                    break
                labelled.append(
                    {
                        "text": text,
                        "topic_labels": "|".join(topics),
                        "quality_score": quality,
                        "difficulty_level": difficulty,
                    }
                )
                count += 1
                console.print(
                    f"[green]✓ Re-labelled[/green] "
                    f"[dim]Topics: {', '.join(topics[:3])}, Quality: {quality}, Difficulty: {difficulty}[/dim]"
                )
        if count > 0:
            Path(out_path).parent.mkdir(parents=True, exist_ok=True)
            with open(out_path, "w", newline="") as f:
                w = csv.DictWriter(
                    f,
                    fieldnames=[
                        "text",
                        "topic_labels",
                        "quality_score",
                        "difficulty_level",
                    ],
                )
                w.writeheader()
                w.writerows(labelled)
            self._train_data = out_path
            console.print(f"[green]✓ Saved {len(labelled)} rows → {out_path}[/green]")
            tlog.info("LABEL  %d new, %d total → %s", count, len(labelled), out_path)
        else:
            console.print("[dim]No new labels added[/dim]")

    # ── bulk label ──────────────────────────────────────────────
    def bulk_label_flow(self):
        """Fast bulk labelling: show text, single-keystroke quality + difficulty, auto-advance."""
        console.print("\n[bold yellow]Bulk Label Corpus[/bold yellow]")
        console.print(
            "[dim]Fast labelling: rate quality (1-9) and difficulty (e/m/h) per entry.\n"
            "Press 'q' to save and quit. Labels auto-save in batches of 10.[/dim]\n"
        )
        corpus_path = _path("Corpus JSON to label", default=self._corpus_path)
        if corpus_path is None:
            return
        out_path = _path("Output labelled CSV", default=self._train_data)
        if out_path is None:
            return
        try:
            with open(corpus_path) as f:
                data = json.load(f)
            entries = data if isinstance(data, list) else data.get("entries", [])
        except Exception as e:
            console.print(f"[red]✗ Failed to load corpus: {e}[/red]")
            return
        if not entries:
            console.print("[red]✗ No entries[/red]")
            return

        existing = []
        skip_texts = set()
        if Path(out_path).exists():
            with open(out_path) as f:
                reader = csv.DictReader(f)
                for row in reader:
                    existing.append(row)
                    skip_texts.add(row.get("text", "")[:100])
        unlabelled = [e for e in entries if e.get("text", "")[:100] not in skip_texts]
        if not unlabelled:
            console.print("[green]All entries already labelled[/green]")
            return
        console.print(f"[cyan]{len(unlabelled)} entries to label[/cyan]")

        diff_map = {"e": "easy", "m": "medium", "h": "hard"}
        labelled = list(existing)
        count = 0

        for i, entry in enumerate(unlabelled):
            text = entry.get("text", "")
            topics_hint = entry.get("topic", entry.get("topics", ""))
            if isinstance(topics_hint, list):
                topics_hint = ", ".join(topics_hint)
            console.print(
                Panel(
                    text[:800] + ("..." if len(text) > 800 else ""),
                    title=f"[{i+1}/{len(unlabelled)}] Topics: {topics_hint}",
                    box=box.ROUNDED,
                    border_style="cyan",
                )
            )
            quality = _text("Quality (1-9, q to quit)", default="7")
            if quality is None or quality.lower() == "q":
                break
            try:
                q = int(quality)
                if q < 1 or q > 9:
                    console.print("[red]Must be 1-9[/red]")
                    continue
            except ValueError:
                console.print("[red]Must be a number 1-9[/red]")
                continue
            diff = _text("Difficulty (e/m/h)", default="m")
            if diff is None or diff.lower() == "q":
                break
            if diff.lower() not in diff_map:
                console.print("[red]Must be e, m, or h[/red]")
                continue
            labelled.append(
                {
                    "text": text,
                    "topic_labels": topics_hint
                    if isinstance(topics_hint, str)
                    else "|".join(topics_hint),
                    "quality_score": str(q),
                    "difficulty_level": diff_map[diff.lower()],
                }
            )
            count += 1
            console.print(f"[green]✓ {count} labelled[/green]")

            # Auto-save every 10
            if count % 10 == 0:
                self._write_labels(out_path, labelled)
                console.print(f"[dim]Auto-saved {len(labelled)} rows[/dim]")

        if count > 0:
            self._write_labels(out_path, labelled)
            console.print(f"[green]✓ Saved {len(labelled)} rows → {out_path}[/green]")
            tlog.info("BULK_LABEL  %d new, %d total", count, len(labelled))
        else:
            console.print("[dim]No new labels added[/dim]")

    def _write_labels(self, out_path, labelled):
        Path(out_path).parent.mkdir(parents=True, exist_ok=True)
        with open(out_path, "w", newline="") as f:
            w = csv.DictWriter(
                f,
                fieldnames=[
                    "text",
                    "topic_labels",
                    "quality_score",
                    "difficulty_level",
                ],
            )
            w.writeheader()
            w.writerows(labelled)

    # ── batch generate ──────────────────────────────────────────
    def batch_generate_flow(self):
        """Generate multiple hypotheticals in batch."""
        console.print("\n[bold yellow]Batch Generate Hypotheticals[/bold yellow]")
        count = _validated_text(
            "Number to generate (2-10)", default="3", validate=_validate_number(2, 10)
        )
        if count is None:
            return
        count = int(count)
        strategy = _select_quit(
            "Topic spread",
            choices=[
                Choice("Random — pick topics randomly", value="random"),
                Choice("Specified — choose topics", value="specified"),
                Choice("Balanced — cycle through all topics", value="balanced"),
            ],
        )
        if strategy is None:
            return
        topics_list: List[str] = []
        if strategy == "specified":
            selected = _checkbox("Select topics", choices=_topic_choices())
            if not selected:
                return
            topics_list = selected
        elif strategy == "balanced":
            topics_list = list(TOPICS)
        elif strategy == "random":
            import random

            topics_list = random.sample(TOPICS, min(count, len(TOPICS)))

        provider = _select_quit("Provider", choices=PROVIDER_CHOICES)
        if provider is None:
            return
        # Auto-check provider health
        with console.status(f"[dim]Checking {provider}...", spinner="dots"):
            healthy = self._ping_provider(provider)
        if not healthy:
            console.print(f"[yellow]⚠ {provider} unreachable — check connection or API key[/yellow]")
            if not _confirm(f"Continue with {provider} anyway?", default=False):
                return
        model_name = _select_quit("Model", choices=_model_choices(provider))
        if model_name is None:
            return
        if model_name == "__custom__":
            model_name = _text("Custom model name", default="")
            if model_name is None:
                return
        complexity = _select_quit("Complexity", choices=COMPLEXITY_CHOICES)
        if complexity is None:
            return
        parties = _select_quit("Parties", choices=PARTIES_CHOICES)
        if parties is None:
            return
        method = _select_quit("Method", choices=METHOD_CHOICES)
        if method is None:
            return
        if not _confirm(f"Generate {count} hypotheticals?", default=True):
            return

        from ..services.hypothetical_service import (
            GenerationRequest,
            hypothetical_service,
        )

        results = []
        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            BarColumn(),
            TextColumn("{task.completed}/{task.total}"),
            TimeElapsedColumn(),
            console=console,
        ) as progress:
            task = progress.add_task("[cyan]Generating batch", total=count)
            for i in range(count):
                # Pick topic based on strategy
                if strategy == "balanced":
                    topic = topics_list[i % len(topics_list)]
                elif strategy == "specified":
                    topic = topics_list[i % len(topics_list)]
                else:
                    import random

                    topic = random.choice(topics_list)
                progress.update(
                    task,
                    description=f"[cyan]Generating {i+1}/{count}: {topic}",
                )
                try:

                    async def _gen(t=topic):
                        req = GenerationRequest(
                            topics=[t],
                            number_parties=max(2, min(5, int(parties))),
                            complexity_level=str(complexity),
                            method=method,
                            provider=provider,
                            model=model_name or None,
                        )
                        return await hypothetical_service.generate_hypothetical(req)

                    resp = _run_async(_gen(topic))
                    score = resp.validation_results.get("quality_score", 0.0)
                    results.append(
                        {
                            "topic": topic,
                            "score": score,
                            "words": len(resp.hypothetical.split()),
                        }
                    )
                except Exception as e:
                    results.append(
                        {"topic": topic, "score": 0.0, "words": 0, "error": str(e)}
                    )
                progress.advance(task)

        # Summary table
        st = Table(title=f"Batch Results ({count} hypotheticals)", box=box.ROUNDED)
        st.add_column("#", style="cyan", width=3)
        st.add_column("Topic", style="yellow")
        st.add_column("Score", style="green")
        st.add_column("Words", style="dim")
        st.add_column("Status", style="cyan")
        for i, r in enumerate(results, 1):
            status = (
                "[green]OK[/green]" if r.get("score", 0) >= 7.0 else "[red]LOW[/red]"
            )
            if "error" in r:
                status = f"[red]ERR: {r['error'][:30]}[/red]"
            st.add_row(str(i), r["topic"], f"{r['score']:.1f}", str(r["words"]), status)
        console.print(st)
        tlog.info("BATCH  generated %d hypotheticals", count)

    # ── export ─────────────────────────────────────────────────
    def export_flow(self):
        """Export a generated hypothetical + analysis to DOCX (and optionally PDF)."""
        console.print("\n[bold yellow]Export Hypothetical to DOCX/PDF[/bold yellow]")
        console.print(
            "[dim]Export a generated hypothetical to a document.\n"
            "Requires python-docx (pip install python-docx).\n"
            "PDF conversion requires docx2pdf (optional).[/dim]\n"
        )
        history = self._load_history()
        hypo_text = ""
        analysis_text = ""
        model_answer = ""

        # Offer options: export from history or paste manually
        source_choices = []
        if history:
            source_choices.append(Choice("Export last generated", value="last"))
            source_choices.append(Choice("Select from history", value="history"))
        source_choices.append(Choice("Paste text manually", value="paste"))
        source_choices.append(Choice("Back", value="back"))

        source = _select_quit("Export source", choices=source_choices)
        if source is None or source == "back":
            return

        if source == "last" and history:
            rec = history[-1]
            hypo_text = rec.get("hypothetical", "")
            analysis_text = rec.get("analysis", "")
            model_answer = rec.get("model_answer", "")
            topic = rec.get("config", {}).get("topic", "hypothetical")
            console.print(f"[green]✓ Loaded last generated ({topic})[/green]")
        elif source == "history" and history:
            # Show history selection
            ht = Table(title=f"Select from History ({len(history)} records)", box=box.ROUNDED)
            ht.add_column("#", style="cyan", width=4)
            ht.add_column("Time", style="dim", width=16)
            ht.add_column("Topic", style="yellow", width=18)
            ht.add_column("Preview", style="dim")
            display_records = history[-20:]
            for i, r in enumerate(display_records, 1):
                ts = r.get("timestamp", "")[:16]
                cfg = r.get("config", {})
                topic = cfg.get("topic", str(cfg.get("topics", "")))
                text = r.get("hypothetical", "")[:60]
                ht.add_row(str(i), ts, topic, text.replace("\n", " "))
            console.print(ht)
            idx = _text(f"Select record # (1-{len(display_records)})", default="1")
            if not idx:
                return
            try:
                rec = display_records[int(idx) - 1]
                hypo_text = rec.get("hypothetical", "")
                analysis_text = rec.get("analysis", "")
                model_answer = rec.get("model_answer", "")
                console.print(f"[green]✓ Loaded record #{idx}[/green]")
            except (ValueError, IndexError):
                console.print("[red]Invalid selection[/red]")
                return
        elif source == "paste":
            hypo_text = _text("Paste hypothetical text (or path to .txt file)", default="")
            if not hypo_text:
                return
            # Check if it's a file path
            if Path(hypo_text).exists() and Path(hypo_text).is_file():
                hypo_text = Path(hypo_text).read_text(encoding="utf-8")
            analysis_text = _text(
                "Paste analysis text (optional, Enter to skip)", default=""
            )
            model_answer = _text("Paste model answer (optional, Enter to skip)", default="")

        if not hypo_text:
            console.print("[red]✗ No hypothetical text to export[/red]")
            return

        out_path = _path("Output DOCX path", default="exports/hypothetical.docx")
        if not out_path:
            return
        also_pdf = _confirm("Also export PDF?", default=False)
        try:
            import docx
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt

            doc = docx.Document()
            style = doc.styles["Normal"]
            font = style.font
            font.name = "Times New Roman"
            font.size = Pt(12)

            doc.add_heading("Legal Hypothetical", level=1)
            for i, para in enumerate(hypo_text.split("\n\n")):
                p = doc.add_paragraph(para.strip())
                if i == 0:
                    p.paragraph_format.space_before = Pt(12)

            if analysis_text:
                doc.add_heading("Legal Analysis", level=1)
                for para in analysis_text.split("\n\n"):
                    doc.add_paragraph(para.strip())

            if model_answer:
                doc.add_heading("Model Answer", level=1)
                for para in model_answer.split("\n\n"):
                    doc.add_paragraph(para.strip())

            # Footer with metadata
            from datetime import datetime

            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = (
                f"Generated by Jikai | {datetime.now().strftime('%Y-%m-%d %H:%M')}"
            )
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            Path(out_path).parent.mkdir(parents=True, exist_ok=True)
            doc.save(out_path)
            console.print(f"[green]✓ DOCX saved → {out_path}[/green]")

            if also_pdf:
                try:
                    from docx2pdf import convert

                    pdf_path = str(Path(out_path).with_suffix(".pdf"))
                    convert(out_path, pdf_path)
                    console.print(f"[green]✓ PDF saved → {pdf_path}[/green]")
                except ImportError:
                    console.print(
                        "[yellow]docx2pdf not installed. Install with: pip install docx2pdf[/yellow]"
                    )
            tlog.info("EXPORT  %s", out_path)
        except ImportError:
            console.print(
                "[red]✗ python-docx not installed. Run: pip install python-docx[/red]"
            )
        except Exception as e:
            console.print(f"[red]✗ Export failed: {e}[/red]")
            tlog.info("ERROR  export: %s", e)

    # ── import cases ─────────────────────────────────────────────
    def import_cases_flow(self):
        """Import SG case summaries from PDF/DOCX into corpus/cases/ as structured JSON."""
        console.print("\n[bold yellow]Import SG Case Summaries[/bold yellow]")
        console.print(
            "[dim]Import PDF/DOCX files containing SG case summaries.\n"
            "Text will be extracted and stored as structured JSON in corpus/cases/.[/dim]\n"
        )
        file_path = _path("Path to case file (PDF/DOCX)", default="")
        if not file_path:
            return
        p = Path(file_path)
        if not p.exists():
            console.print(f"[red]✗ File not found: {file_path}[/red]")
            return
        if p.suffix.lower() not in (".pdf", ".docx", ".txt"):
            console.print("[red]✗ Unsupported format. Use PDF, DOCX, or TXT.[/red]")
            return
        try:
            from ..services.corpus_preprocessor import extract_text, normalize_text

            with console.status("[bold green]Extracting text...", spinner="dots"):
                raw_text = normalize_text(extract_text(p))
            if not raw_text or len(raw_text) < 50:
                console.print("[red]✗ No usable text extracted from file[/red]")
                return
            console.print(
                Panel(
                    raw_text[:1500] + ("..." if len(raw_text) > 1500 else ""),
                    title=f"Extracted from {p.name}",
                    box=box.ROUNDED,
                    border_style="green",
                )
            )
            case_name = _text("Case name", default=p.stem.replace("_", " ").title())
            if case_name is None:
                return
            citation = _text("Citation (e.g. [2020] SGCA 1)", default="")
            if citation is None:
                return
            summary = _text(
                "Brief summary (or press Enter to use full text)", default=""
            )
            topics = _checkbox("Related topics", choices=_topic_choices())
            if topics is None:
                return
            case_entry = {
                "case_name": case_name,
                "citation": citation,
                "summary": summary if summary else raw_text,
                "full_text": raw_text,
                "topics": topics,
                "source_file": str(p),
            }
            cases_dir = Path("corpus/cases")
            cases_dir.mkdir(parents=True, exist_ok=True)
            safe_name = re.sub(r"[^\w\-]", "_", case_name.lower())[:60]
            out_path = cases_dir / f"{safe_name}.json"
            with open(out_path, "w", encoding="utf-8") as f:
                json.dump(case_entry, f, indent=2, ensure_ascii=False)
            console.print(f"[green]✓ Case saved → {out_path}[/green]")
            tlog.info("IMPORT_CASE  %s → %s", case_name, out_path)
        except Exception as e:
            console.print(f"[red]✗ Import failed: {e}[/red]")
            tlog.info("ERROR  import_cases: %s", e)

    # ── generate ────────────────────────────────────────────────
    def _load_state(self) -> Dict:
        """Load persisted state from .jikai_state."""
        try:
            with open(".jikai_state", "r") as f:
                return json.load(f)
        except json.JSONDecodeError:
            import shutil

            shutil.copy2(".jikai_state", ".jikai_state.bak")
            console.print(
                "[yellow]⚠ .jikai_state corrupted, backed up to .jikai_state.bak[/yellow]"
            )
            return {}
        except Exception:
            return {}

    def _save_state(self, state: Dict):
        """Save state to .jikai_state."""
        tmp = ".jikai_state.tmp"
        with open(tmp, "w") as f:
            json.dump(state, f, indent=2)
        os.rename(tmp, ".jikai_state")  # atomic write

    def generate_flow(self):
        while True:
            console.print("\n[bold yellow]Generate Hypothetical[/bold yellow]")
            mode = _select_quit(
                "Mode",
                choices=[
                    Choice("Quick generate — topic only, use defaults", value="quick"),
                    Choice("Custom generate — full config", value="custom"),
                ],
            )
            if mode is None:
                return

            state = self._load_state()
            defaults = state.get("last_config", {})

            if mode == "quick":
                # Show defaults first before committing
                provider = defaults.get("provider", "ollama")
                model_name = defaults.get("model", None)
                temperature = defaults.get("temperature", 0.7)
                complexity = defaults.get("complexity", "3")
                parties = defaults.get("parties", "2")
                method = defaults.get("method", "pure_llm")
                red_herrings = False

                dt = Table(title="Quick Generate Defaults", box=box.SIMPLE)
                dt.add_column("Setting", style="cyan")
                dt.add_column("Value", style="yellow")
                dt.add_row("Provider", provider)
                dt.add_row("Model", model_name or "default")
                dt.add_row("Temperature", str(temperature))
                dt.add_row("Complexity", complexity)
                dt.add_row("Parties", parties)
                dt.add_row("Method", method)
                console.print(dt)

                if not _confirm("Use these settings?", default=True):
                    console.print("[dim]Switching to custom mode...[/dim]")
                    mode = "custom"
                else:
                    topic = _select_quit("Topic", choices=_topic_choices())
                    if topic is None:
                        return

            if mode == "custom":
                topic = _select_quit("Topic", choices=_topic_choices())
                if topic is None:
                    return
                provider = _select_quit("Provider", choices=PROVIDER_CHOICES)
                if provider is None:
                    return
                # Auto-check provider health
                with console.status(f"[dim]Checking {provider}...", spinner="dots"):
                    healthy = self._ping_provider(provider)
                if not healthy:
                    console.print(f"[yellow]⚠ {provider} unreachable — check connection or API key[/yellow]")
                    if not _confirm(f"Continue with {provider} anyway?", default=False):
                        return
                model_name = _select_quit("Model", choices=_model_choices(provider))
                if model_name is None:
                    return
                if model_name == "__custom__":
                    model_name = _text("Custom model name", default="")
                    if model_name is None:
                        return
                temperature = _select_quit(
                    "Temperature",
                    choices=[
                        Choice(f"{v/10:.1f}", value=str(v / 10)) for v in range(1, 16)
                    ],
                )
                if temperature is None:
                    return
                temperature = float(temperature)
                complexity = _select_quit("Complexity", choices=COMPLEXITY_CHOICES)
                if complexity is None:
                    return
                parties = _select_quit("Parties", choices=PARTIES_CHOICES)
                if parties is None:
                    return
                method = _select_quit("Method", choices=METHOD_CHOICES)
                if method is None:
                    return
                red_herrings = _confirm("Include red herrings?", default=False)

            cfg = Table(box=box.SIMPLE, title="Generation Config")
            cfg.add_column("Parameter", style="cyan")
            cfg.add_column("Value", style="yellow")
            cfg.add_row("Topic", topic)
            cfg.add_row("Provider", provider)
            cfg.add_row("Model", str(model_name) or "(default)")
            cfg.add_row("Temperature", f"{float(temperature):.1f}")
            cfg.add_row("Complexity", str(complexity))
            cfg.add_row("Parties", str(parties))
            cfg.add_row("Method", method)
            cfg.add_row("Red Herrings", "Yes" if red_herrings else "No")
            console.print(cfg)
            if not _confirm("Proceed with generation?", default=True):
                return
            tlog.info(
                "GENERATE  topic=%s provider=%s model=%s temp=%.1f "
                "complexity=%s parties=%s method=%s",
                topic,
                provider,
                model_name,
                temperature,
                complexity,
                parties,
                method,
            )
            self._do_generate(
                GenerationConfig(
                    topic=topic,
                    provider=provider,
                    model=model_name or None,
                    complexity=int(complexity),
                    parties=int(parties),
                    method=method,
                    temperature=float(temperature),
                    red_herrings=red_herrings,
                )
            )
            # Persist last-used config for quick-generate
            state = self._load_state()
            state["last_config"] = {
                "provider": provider,
                "model": model_name,
                "temperature": float(temperature),
                "complexity": str(complexity),
                "parties": str(parties),
                "method": method,
            }
            self._save_state(state)
            if not _confirm("Generate another?", default=False):
                return

    def _do_generate(self, cfg: "GenerationConfig"):
        topic, provider, model = cfg.topic, cfg.provider, cfg.model
        complexity, parties, method = cfg.complexity, cfg.parties, cfg.method
        temperature, red_herrings = cfg.temperature, cfg.red_herrings
        max_retries = 3
        partial_result = ""
        try:
            from ..services.llm_service import LLMRequest, llm_service
            from ..services.validation_service import validation_service

            red_herring_instr = ""
            if red_herrings:
                red_herring_instr = (
                    " Include 1-2 legally irrelevant but plausible facts as red herrings."
                    " The red herrings should not dominate the scenario."
                )
            request = LLMRequest(
                prompt=(
                    f"Generate a Singapore tort law hypothetical about "
                    f"{topic} with {parties} parties at complexity "
                    f"{complexity}/5.{red_herring_instr}"
                ),
                temperature=temperature,
                stream=True,
            )

            async def _stream():
                nonlocal partial_result
                chunks: List[str] = []
                with Live(
                    Text("Generating..."), console=console, refresh_per_second=4
                ) as live:
                    async for chunk in llm_service.stream_generate(
                        request, provider=provider, model=model
                    ):
                        chunks.append(chunk)
                        partial_result = "".join(chunks)
                        live.update(Text(partial_result))
                return partial_result

            try:
                console.print("[dim]Attempt 1/{0}[/dim]".format(max_retries))
                result = _run_async(_stream())
                if result:
                    vr = validation_service.validate_hypothetical(
                        text=result,
                        required_topics=[topic],
                        expected_parties=parties,
                    )
                    score = vr.get("overall_score", 0.0)
                    passed = vr.get("passed", False)
                    console.print(
                        Panel(
                            result,
                            title="Generated Hypothetical",
                            box=box.ROUNDED,
                            border_style="green",
                        )
                    )
                    self._show_validation(vr)
                    if passed or score >= 7.0:
                        tlog.info(
                            "GENERATE  stream success, len=%d score=%.1f",
                            len(result),
                            score,
                        )
                        self._save_to_history(
                            {
                                "config": {
                                    "topic": topic,
                                    "provider": provider,
                                    "model": model,
                                    "complexity": complexity,
                                    "parties": parties,
                                    "method": method,
                                },
                                "hypothetical": result,
                                "validation_score": score,
                            }
                        )
                        self._offer_model_answer(
                            result, topic, provider, model, temperature
                        )
                        self._offer_export_now()
                        return
                    retry_reason = self._build_retry_reason(vr, score)
                    console.print(f"[yellow]{retry_reason}[/yellow]")
            except Exception as stream_err:
                tlog.info("GENERATE  stream failed: %s, trying full", stream_err)
                console.print(
                    "[yellow]Stream unavailable "
                    f"({stream_err}), using full generation...[/yellow]"
                )

            from ..services.hypothetical_service import (
                GenerationRequest,
                hypothetical_service,
            )

            feedback = ""
            for attempt in range(1, max_retries + 1):
                console.print(
                    f"[dim]Attempt {attempt}/{max_retries} (full generation)[/dim]"
                )

                async def _full(extra_feedback=feedback):
                    prefs = {"temperature": temperature, "red_herrings": red_herrings}
                    if extra_feedback:
                        prefs["feedback"] = extra_feedback
                    req = GenerationRequest(
                        topics=[topic],
                        number_parties=max(2, min(5, parties)),
                        complexity_level=str(complexity),
                        method=method,
                        provider=provider,
                        model=model,
                        user_preferences=prefs,
                    )
                    return await hypothetical_service.generate_hypothetical(req)

                with console.status(
                    f"[bold green]Generating hypothetical (attempt {attempt}/{max_retries})...",
                    spinner="dots",
                ):
                    response = _run_async(_full(feedback))

                partial_result = response.hypothetical
                console.print(
                    Panel(
                        response.hypothetical,
                        title=f"Generated Hypothetical (attempt {attempt})",
                        box=box.ROUNDED,
                        border_style="green",
                    )
                )
                if response.analysis:
                    console.print(
                        Panel(
                            response.analysis,
                            title="Legal Analysis",
                            box=box.ROUNDED,
                            border_style="cyan",
                        )
                    )
                vr = response.validation_results
                self._show_validation(vr)
                score = vr.get("quality_score", 0.0)
                passed = vr.get("passed", False)
                if passed or score >= 7.0:
                    tlog.info(
                        "GENERATE  full success attempt=%d score=%.1f",
                        attempt,
                        score,
                    )
                    self._save_to_history(
                        {
                            "config": {
                                "topic": topic,
                                "provider": provider,
                                "model": model,
                                "complexity": complexity,
                                "parties": parties,
                                "method": method,
                            },
                            "hypothetical": response.hypothetical,
                            "analysis": response.analysis,
                            "validation_score": score,
                        }
                    )
                    self._offer_model_answer(
                        response.hypothetical, topic, provider, model, temperature
                    )
                    self._offer_export_now()
                    return

                if attempt < max_retries:
                    # Build feedback for next attempt
                    checks = vr.get(
                        "checks", vr.get("adherence_check", {}).get("checks", {})
                    )
                    issues = []
                    if isinstance(checks, dict):
                        pc = checks.get("party_count", {})
                        if not pc.get("passed", True):
                            found = pc.get("found", "?")
                            issues.append(
                                f"Include exactly {parties} distinct named parties (found {found})."
                            )
                        ti = checks.get("topic_inclusion", {})
                        if not ti.get("passed", True):
                            missing = ti.get("topics_missing", [])
                            covered = ti.get("topics_found", [])
                            issues.append(
                                f"Must address these topics: {', '.join(missing)}."
                                + (f" (covered: {', '.join(covered)})" if covered else "")
                            )
                        wc = checks.get("word_count", {})
                        if not wc.get("passed", True):
                            actual = wc.get("count", wc.get("word_count", "?"))
                            issues.append(f"Aim for 800-1500 words (currently {actual}).")
                        if not checks.get("singapore_context", {}).get("passed", True):
                            issues.append("Set the scenario explicitly in Singapore.")
                    feedback = "Previous attempt scored {:.1f}/10. Fix these issues: {}".format(
                        score,
                        " ".join(issues) if issues else "Improve overall quality.",
                    )
                    retry_reason = self._build_retry_reason(vr, score)
                    console.print(f"[yellow]{retry_reason}[/yellow]")

            console.print(
                "[yellow]Max retries reached. Showing best result above.[/yellow]"
            )
            tlog.info("GENERATE  max retries reached, score=%.1f", score)
        except ImportError as e:
            console.print(f"[red]✗ Missing dependency: {e}[/red]")
            console.print(
                "[yellow]Run: pip install chromadb sentence-transformers[/yellow]"
            )
            tlog.info("ERROR  import: %s", e)
        except KeyboardInterrupt:
            word_count = len(partial_result.split()) if partial_result else 0
            if partial_result and word_count > 20:
                console.print(
                    f"\n[yellow]Generation cancelled. Partial result ({word_count} words):[/yellow]"
                )
                console.print(
                    Panel(
                        partial_result,
                        title="Partial Hypothetical",
                        box=box.ROUNDED,
                        border_style="yellow",
                    )
                )
            else:
                console.print("\n[yellow]Generation cancelled.[/yellow]")
            tlog.info("GENERATE  cancelled by user")
        except Exception as e:
            self._handle_generation_error(e)
            tlog.info("ERROR  generation: %s", e)

    def _handle_generation_error(self, e: Exception):
        """Pattern-match common errors and provide specific recovery messages."""
        err_str = str(e).lower()

        # API key errors
        if "api key" in err_str or "apikey" in err_str or "authentication" in err_str:
            if "openai" in err_str:
                console.print("[red]✗ No API key set for OpenAI. Go to More → Settings to add one.[/red]")
            elif "anthropic" in err_str:
                console.print("[red]✗ No API key set for Anthropic. Go to More → Settings to add one.[/red]")
            elif "google" in err_str:
                console.print("[red]✗ No API key set for Google. Go to More → Settings to add one.[/red]")
            else:
                console.print("[red]✗ API key missing or invalid. Go to More → Settings to check your keys.[/red]")
        # Model not found
        elif "model" in err_str and ("not found" in err_str or "does not exist" in err_str or "not available" in err_str):
            console.print("[red]✗ Model not available. Run More → Providers → Check Health to see available models.[/red]")
        # Connection/network errors — attempt to auto-restart Ollama
        elif "timeout" in err_str or "timed out" in err_str or "connection" in err_str or "refused" in err_str or "unreachable" in err_str or "all connection attempts failed" in err_str:
            console.print("[red]✗ Cannot connect to provider.[/red]")
            provider_name = self._configured_provider()
            if provider_name == "ollama":
                host = self._ollama_host_from_env()
                started = self._start_ollama(host)
                if started:
                    console.print(
                        "[dim]Ollama is now running — press Generate again to retry.[/dim]"
                    )
                else:
                    console.print(
                        "[dim]Could not start Ollama automatically. "
                        "Run 'ollama serve' in a terminal, then try again.[/dim]"
                    )
            else:
                console.print("[dim]Check that your provider service is running and accessible.[/dim]")
        # Rate limiting
        elif "rate limit" in err_str or "too many requests" in err_str:
            console.print("[red]✗ Rate limited by provider. Wait a moment and try again.[/red]")
        # Generic fallback
        else:
            console.print(f"[red]✗ Generation failed: {e}[/red]")
            console.print("[dim]Tip: check provider health via More → Providers[/dim]")

    def _offer_model_answer(
        self, hypothetical_text, topic, provider, model, temperature
    ):
        """Offer to generate a model answer (IRAC analysis) for the hypothetical."""
        if not _confirm("Generate model answer?", default=False):
            return
        try:
            from ..services.llm_service import LLMRequest, llm_service

            prompt = (
                "You are an expert Singapore tort law tutor. Given the following hypothetical, "
                "produce a model answer consisting of:\n"
                "1. ISSUE-SPOTTING CHECKLIST: List every legal issue present.\n"
                "2. STRUCTURED IRAC ANALYSIS: For each issue, provide:\n"
                "   - Issue: State the legal question\n"
                "   - Rule: State the relevant legal principle(s) and SG case law\n"
                "   - Application: Apply the rule to the facts\n"
                "   - Conclusion: State the likely outcome\n\n"
                f"HYPOTHETICAL:\n{hypothetical_text}"
            )
            req = LLMRequest(
                prompt=prompt,
                temperature=max(0.3, temperature - 0.2),
                max_tokens=3000,
            )

            async def _gen():
                return await llm_service.generate(req, provider=provider, model=model)

            with console.status(
                "[bold green]Generating model answer...", spinner="dots"
            ):
                resp = _run_async(_gen())
            console.print(
                Panel(
                    resp.content,
                    title="Model Answer",
                    box=box.ROUNDED,
                    border_style="magenta",
                )
            )
            tlog.info("MODEL_ANSWER  generated, len=%d", len(resp.content))
        except Exception as e:
            console.print(f"[red]✗ Model answer generation failed: {e}[/red]")

    def _show_validation(self, vr):
        """Display validation results table with detailed check results."""
        if not vr:
            return
        vt = Table(title="Validation Results", box=box.ROUNDED)
        vt.add_column("Check", style="cyan")
        vt.add_column("Result", style="yellow")

        # Extract detailed checks
        checks = vr.get("checks", vr.get("adherence_check", {}).get("checks", {}))
        if isinstance(checks, dict):
            # Party Count
            pc = checks.get("party_count", {})
            if pc:
                passed = pc.get("passed", True)
                found = pc.get("found", "?")
                expected = pc.get("expected", "?")
                icon = "[green]✓[/green]" if passed else "[red]✗[/red]"
                vt.add_row("Party Count", f"{icon} Found {found} (expected {expected})")

            # Topic Coverage
            ti = checks.get("topic_inclusion", {})
            if ti:
                passed = ti.get("passed", True)
                icon = "[green]✓[/green]" if passed else "[red]✗[/red]"
                if passed:
                    vt.add_row("Topic Coverage", f"{icon} Covered")
                else:
                    missing = ti.get("missing", ti.get("topics_missing", []))
                    if missing:
                        vt.add_row("Topic Coverage", f"{icon} Missing: {', '.join(missing)}")
                    else:
                        vt.add_row("Topic Coverage", f"{icon} Incomplete")

            # Word Count
            wc = checks.get("word_count", {})
            if wc:
                passed = wc.get("passed", True)
                count = wc.get("count", wc.get("word_count", "?"))
                icon = "[green]✓[/green]" if passed else "[red]✗[/red]"
                vt.add_row("Word Count", f"{icon} {count} words (target: 800-1500)")

            # Singapore Context
            sc = checks.get("singapore_context", {})
            if sc:
                passed = sc.get("passed", True)
                icon = "[green]✓[/green]" if passed else "[red]✗[/red]"
                vt.add_row("Singapore Context", f"{icon} {'Present' if passed else 'Missing'}")

        # Quality Score
        score = vr.get("quality_score", vr.get("overall_score", "N/A"))
        if isinstance(score, (int, float)):
            vt.add_row("Quality Score", f"{score:.1f} / 10")
        else:
            vt.add_row("Quality Score", str(score))

        # Overall
        passed = vr.get("passed", False)
        vt.add_row(
            "Overall",
            "[green]✓ Passed[/green]" if passed else "[red]✗ Failed[/red]",
        )
        console.print(vt)

    def _build_retry_reason(self, vr, score):
        """Build a detailed retry reason message from validation results."""
        checks = vr.get("checks", vr.get("adherence_check", {}).get("checks", {}))
        failed_checks = []
        if isinstance(checks, dict):
            pc = checks.get("party_count", {})
            if not pc.get("passed", True):
                found = pc.get("found", "?")
                expected = pc.get("expected", "?")
                failed_checks.append(f"party count ({found}/{expected})")
            ti = checks.get("topic_inclusion", {})
            if not ti.get("passed", True):
                missing = ti.get("missing", [])
                if missing:
                    failed_checks.append(f"topic coverage (missing: {', '.join(missing)})")
                else:
                    failed_checks.append("topic coverage")
            wc = checks.get("word_count", {})
            if not wc.get("passed", True):
                actual = wc.get("count", wc.get("word_count", "?"))
                failed_checks.append(f"word count ({actual} words)")
            sc = checks.get("singapore_context", {})
            if not sc.get("passed", True):
                failed_checks.append("Singapore context")
        if failed_checks:
            return f"Score {score:.1f}/10 — minimum is 7.0, retrying (failed: {', '.join(failed_checks)})..."
        return f"Score {score:.1f}/10 — minimum is 7.0, retrying with corrective feedback..."

    def _offer_export_now(self):
        """Offer to export the last generated hypothetical immediately."""
        if not _confirm("Export now?", default=False):
            return
        self.export_flow()

    def _save_to_history(self, record: Dict):
        """Append a generation record to data/history.json."""
        from datetime import datetime

        record["timestamp"] = datetime.now().isoformat()
        hp = Path(HISTORY_PATH)
        hp.parent.mkdir(parents=True, exist_ok=True)
        history = []
        if hp.exists():
            try:
                with open(hp, "r") as f:
                    history = json.load(f)
            except Exception:
                history = []
        history.append(record)
        record_num = len(history)
        if len(history) > 500:  # cap history size
            history = history[-500:]
        tmp = hp.with_suffix(".json.tmp")
        with open(tmp, "w") as f:
            json.dump(history, f, indent=2, ensure_ascii=False)
        os.rename(tmp, hp)  # atomic write
        console.print(f"[green]✓ Saved to history (#{record_num})[/green]")

    def _load_history(self) -> List[Dict]:
        """Load generation history from data/history.json."""
        hp = Path(HISTORY_PATH)
        if not hp.exists():
            return []
        try:
            with open(hp, "r") as f:
                return json.load(f)
        except json.JSONDecodeError:
            import shutil

            bak = hp.with_suffix(".json.bak")
            shutil.copy2(hp, bak)
            console.print(
                f"[yellow]⚠ history.json corrupted, backed up to {bak}[/yellow]"
            )
            return []
        except Exception:
            return []

    # ── history ────────────────────────────────────────────────
    def history_flow(self):
        """Browse generation history with search/filter/recall."""
        console.print("\n[bold yellow]Generation History[/bold yellow]")
        history = self._load_history()
        if not history:
            console.print(
                "[dim]No generation history yet. Generate a hypothetical first.[/dim]"
            )
            return
        while True:
            c = _select_quit(
                f"History ({len(history)} records)",
                choices=[
                    Choice("Browse all", value="browse"),
                    Choice("Search by keyword", value="search"),
                    Choice("Filter by topic", value="filter"),
                ],
            )
            if c is None:
                return
            if c == "search":
                term = _text("Search keyword")
                if not term:
                    continue
                filtered = [h for h in history if term.lower() in json.dumps(h).lower()]
                console.print(f"[green]{len(filtered)} matches[/green]")
                self._display_history(filtered)
            elif c == "filter":
                topic = _select_quit("Filter by topic", choices=_topic_choices())
                if not topic:
                    continue
                filtered = [
                    h
                    for h in history
                    if topic in h.get("config", {}).get("topic", "")
                    or topic in str(h.get("config", {}).get("topics", []))
                ]
                console.print(f"[green]{len(filtered)} matches[/green]")
                self._display_history(filtered)
            elif c == "browse":
                self._display_history(history)

    def _display_history(self, records: List[Dict]):
        """Display history records in a table."""
        if not records:
            console.print("[dim]No records to display[/dim]")
            return
        ht = Table(title=f"History ({len(records)} records)", box=box.ROUNDED)
        ht.add_column("#", style="cyan", width=4)
        ht.add_column("Time", style="dim", width=16)
        ht.add_column("Topic", style="yellow", width=18)
        ht.add_column("Score", style="green", width=6)
        ht.add_column("Preview", style="dim")
        for i, r in enumerate(records[-20:], 1):
            ts = r.get("timestamp", "")[:16]
            cfg = r.get("config", {})
            topic = cfg.get("topic", str(cfg.get("topics", "")))
            score = str(r.get("validation_score", r.get("score", "N/A")))
            text = r.get("hypothetical", "")[:60]
            ht.add_row(str(i), ts, topic, score, text.replace("\n", " "))
        console.print(ht)
        # Option to view full record
        idx = _text(
            f"View record # (1-{min(20, len(records))}, Enter to skip)", default=""
        )
        if idx:
            try:
                rec = records[-20:][int(idx) - 1]
                console.print(
                    Panel(
                        rec.get("hypothetical", "No text"),
                        title=f"Hypothetical — {rec.get('config', {}).get('topic', '')}",
                        box=box.ROUNDED,
                        border_style="green",
                    )
                )
                if rec.get("analysis"):
                    console.print(
                        Panel(
                            rec["analysis"],
                            title="Analysis",
                            box=box.ROUNDED,
                            border_style="cyan",
                        )
                    )
                if rec.get("model_answer"):
                    console.print(
                        Panel(
                            rec["model_answer"],
                            title="Model Answer",
                            box=box.ROUNDED,
                            border_style="magenta",
                        )
                    )
                self._offer_variation(rec)
            except (ValueError, IndexError):
                console.print("[red]Invalid selection[/red]")

    # ── stats dashboard ─────────────────────────────────────────
    def stats_flow(self):
        """Display stats from history, labelled data, and model metrics."""
        console.print("\n[bold yellow]Stats Dashboard[/bold yellow]")

        # Generation stats from history
        history = self._load_history()
        gt = Table(title="Generation Stats", box=box.ROUNDED)
        gt.add_column("Metric", style="cyan")
        gt.add_column("Value", style="yellow")
        gt.add_row("Total Generations", str(len(history)))
        if history:
            scores = [
                h.get("validation_score", h.get("score", 0))
                for h in history
                if isinstance(h.get("validation_score", h.get("score")), (int, float))
            ]
            avg_score = sum(scores) / len(scores) if scores else 0
            gt.add_row("Avg Quality Score", f"{avg_score:.1f}")
            # Topic frequency
            topic_counts: Dict[str, int] = {}
            for h in history:
                cfg = h.get("config", {})
                t = cfg.get("topic", "")
                if t:
                    topic_counts[t] = topic_counts.get(t, 0) + 1
            if topic_counts:
                top3 = sorted(topic_counts.items(), key=lambda x: -x[1])[:3]
                gt.add_row(
                    "Top Topics",
                    ", ".join(f"{t}({c})" for t, c in top3),
                )
        console.print(gt)

        # Labelled data stats
        if Path(self._train_data).exists():
            with open(self._train_data) as f:
                reader = csv.DictReader(f)
                rows = list(reader)
            lt = Table(title="Labelled Data Stats", box=box.ROUNDED)
            lt.add_column("Metric", style="cyan")
            lt.add_column("Value", style="yellow")
            lt.add_row("Labelled Entries", str(len(rows)))
            if rows:
                q_scores = []
                for r in rows:
                    try:
                        q_scores.append(float(r.get("quality_score", 0)))
                    except (ValueError, TypeError):
                        pass
                if q_scores:
                    lt.add_row("Avg Quality", f"{sum(q_scores)/len(q_scores):.1f}")
                diff_counts: Dict[str, int] = {}
                for r in rows:
                    d = r.get("difficulty_level", "unknown")
                    diff_counts[d] = diff_counts.get(d, 0) + 1
                lt.add_row(
                    "Difficulty Dist",
                    ", ".join(f"{k}: {v}" for k, v in diff_counts.items()),
                )
            console.print(lt)
        else:
            console.print("[dim]No labelled data found[/dim]")

        # Model status
        mt = Table(title="ML Model Status", box=box.ROUNDED)
        mt.add_column("Model", style="cyan")
        mt.add_column("Status", style="yellow")
        md = Path(self._models_dir)
        for name in ("classifier.joblib", "regressor.joblib", "clusterer.joblib"):
            exists = (md / name).exists()
            mt.add_row(
                name.replace(".joblib", "").title(),
                "[green]Trained[/green]" if exists else "[dim]Not trained[/dim]",
            )
        mt.add_row(
            "Embeddings",
            "[green]Ready[/green]"
            if self._embeddings_ready()
            else "[dim]Not indexed[/dim]",
        )
        console.print(mt)

        # Cost estimate from LLM service
        try:
            from ..services.llm_service import llm_service

            cost = llm_service.get_session_cost()
            ct = Table(title="Session Cost", box=box.ROUNDED)
            ct.add_column("Metric", style="cyan")
            ct.add_column("Value", style="yellow")
            ct.add_row("Total Cost (USD)", f"${cost['total_cost_usd']:.4f}")
            ct.add_row("Input Tokens", str(cost["total_input_tokens"]))
            ct.add_row("Output Tokens", str(cost["total_output_tokens"]))
            console.print(ct)
        except Exception:
            pass

    def _offer_variation(self, original_record: Dict):
        """Offer to create a variation of a generated or historical hypothetical."""
        if not _confirm("Create a variation?", default=False):
            return
        cfg = original_record.get("config", {})
        param = _select_quit(
            "Which parameter to vary?",
            choices=[
                Choice("Swap topic", value="topic"),
                Choice("Change party count", value="parties"),
                Choice("Change complexity", value="complexity"),
                Choice("Change difficulty/method", value="method"),
            ],
        )
        if param is None:
            return
        new_cfg = dict(cfg)
        if param == "topic":
            new_topic = _select_quit("New topic", choices=_topic_choices())
            if new_topic is None:
                return
            new_cfg["topic"] = new_topic
        elif param == "parties":
            new_parties = _select_quit("New party count", choices=PARTIES_CHOICES)
            if new_parties is None:
                return
            new_cfg["parties"] = int(new_parties)
        elif param == "complexity":
            new_c = _select_quit("New complexity", choices=COMPLEXITY_CHOICES)
            if new_c is None:
                return
            new_cfg["complexity"] = int(new_c)
        elif param == "method":
            new_m = _select_quit("New method", choices=METHOD_CHOICES)
            if new_m is None:
                return
            new_cfg["method"] = new_m

        # Generate variation using full generation with reference to original
        try:
            from ..services.hypothetical_service import (
                GenerationRequest,
                hypothetical_service,
            )

            original_text = original_record.get("hypothetical", "")
            prefs = {
                "feedback": (
                    f"Create a variation of this hypothetical by changing the {param}. "
                    f"Original scenario (for reference, do NOT copy):\n{original_text[:500]}..."
                )
            }

            async def _var():
                req = GenerationRequest(
                    topics=[new_cfg.get("topic", cfg.get("topic", "negligence"))],
                    number_parties=max(2, min(5, new_cfg.get("parties", 3))),
                    complexity_level=str(new_cfg.get("complexity", "3")),
                    method=new_cfg.get("method", "pure_llm"),
                    provider=new_cfg.get("provider"),
                    model=new_cfg.get("model"),
                    user_preferences=prefs,
                )
                return await hypothetical_service.generate_hypothetical(req)

            with console.status("[bold green]Generating variation...", spinner="dots"):
                resp = _run_async(_var())
            console.print(
                Panel(
                    resp.hypothetical,
                    title=f"Variation (changed: {param})",
                    box=box.ROUNDED,
                    border_style="yellow",
                )
            )
            score = resp.validation_results.get("quality_score", 0.0)
            self._save_to_history(
                {
                    "config": new_cfg,
                    "hypothetical": resp.hypothetical,
                    "analysis": resp.analysis,
                    "validation_score": score,
                    "variation_of": original_record.get("timestamp", ""),
                }
            )
            console.print(f"[green]✓ Variation saved (score: {score:.1f})[/green]")
        except Exception as e:
            console.print(f"[red]✗ Variation failed: {e}[/red]")

    # ── train ───────────────────────────────────────────────────
    def train_flow(self):
        console.print("\n[bold yellow]Train ML Models[/bold yellow]")
        console.print(
            "[dim]All models are optional — generation works without them.\n"
            "Trained models enrich output with quality checks and topic alignment.[/dim]\n"
        )
        info = Table(box=box.SIMPLE, title="Model Reference")
        info.add_column("Model", style="cyan")
        info.add_column("Purpose", style="yellow")
        info.add_column("Required?", style="dim")
        info.add_row("Classifier", "Predicts legal topics in generated text", "No")
        info.add_row("Regressor", "Scores hypothetical quality (0-10)", "No")
        info.add_row("Clusterer", "Groups similar hypotheticals for diversity", "No")
        console.print(info)
        data_path = _path(
            "Training data path (labelled CSV)",
            default=self._train_data,
        )
        if data_path is None:
            return
        models = _checkbox(
            "Select models to train",
            choices=[
                Choice("Classifier — topic prediction", value="cls", checked=True),
                Choice("Regressor — quality scoring", value="reg", checked=True),
                Choice("Clusterer — diversity grouping", value="clu", checked=True),
            ],
        )
        if models is None:
            return
        train_cls = "cls" in models
        train_reg = "reg" in models
        train_clu = "clu" in models

        # Early exit if no models selected
        if not (train_cls or train_reg or train_clu):
            console.print("[yellow]No models selected. Returning to menu.[/yellow]")
            return

        # Defaults for optional parameters
        max_features = "5000"
        test_split = "0.2"
        cls_c = 1.0
        reg_alpha = 1.0
        n_clusters = 5

        # Text vectorization settings (needed for classifier or regressor)
        if train_cls or train_reg:
            console.print("\n[cyan]Text Vectorization Settings[/cyan]")
            max_features = _validated_text(
                "Max features (vocabulary size)",
                default="5000",
                validate=_validate_number(1, 100000),
            )
            if max_features is None:
                return

            test_split = _validated_text(
                "Test split ratio",
                default="0.2",
                validate=_validate_number(0.0, 1.0, is_float=True),
            )
            if test_split is None:
                return

        # Classifier-specific options
        if train_cls:
            console.print("\n[cyan]Classifier Settings[/cyan]")
            cls_c_str = _validated_text(
                "Classifier regularization (C, higher=less regularization)",
                default="1.0",
                validate=_validate_number(0.001, 1000.0, is_float=True),
            )
            if cls_c_str is None:
                return
            cls_c = float(cls_c_str)

        # Regressor-specific options
        if train_reg:
            console.print("\n[cyan]Regressor Settings[/cyan]")
            reg_alpha_str = _validated_text(
                "Regressor regularization (alpha, higher=more regularization)",
                default="1.0",
                validate=_validate_number(0.0001, 1000.0, is_float=True),
            )
            if reg_alpha_str is None:
                return
            reg_alpha = float(reg_alpha_str)

        # Clustering-specific options (only ask if clustering is selected)
        if train_clu:
            console.print("\n[cyan]Clusterer Settings[/cyan]")
            n_clusters_str = _validated_text(
                "Number of clusters",
                default="5",
                validate=_validate_number(2, 100),
            )
            if n_clusters_str is None:
                return
            n_clusters = int(n_clusters_str)

        cfg = Table(box=box.SIMPLE, title="Training Config")
        cfg.add_column("Parameter", style="cyan")
        cfg.add_column("Value", style="yellow")
        cfg.add_row("Data Path", data_path)
        if train_cls or train_reg:
            cfg.add_row("Max Features", max_features)
            cfg.add_row("Test Split", test_split)
        if train_cls:
            cfg.add_row("Classifier", f"Yes (C={cls_c})")
        if train_reg:
            cfg.add_row("Regressor", f"Yes (alpha={reg_alpha})")
        if train_clu:
            cfg.add_row("Clusterer", f"Yes ({n_clusters} clusters)")
        console.print(cfg)
        if not _confirm("Proceed with training?", default=True):
            return
        tlog.info(
            "TRAIN  data=%s cls=%s reg=%s clu=%s",
            data_path,
            train_cls,
            train_reg,
            train_clu,
        )
        self._do_train(
            data_path,
            int(max_features),
            float(test_split),
            train_cls,
            cls_c,
            train_reg,
            reg_alpha,
            train_clu,
            n_clusters,
        )

    def _do_train(self, data_path, max_features, test_split, train_cls, cls_c, train_reg, reg_alpha, train_clu, n_clusters):
        try:
            from ..ml.pipeline import MLPipeline

            pipeline = MLPipeline()
            stages = {
                "data": "Loading data",
                "features": "Extracting TF-IDF features",
                "classifier": "Training classifier — topic prediction",
                "regressor": "Training regressor — quality scoring",
                "clusterer": "Training clusterer — diversity grouping",
                "saving": "Saving models to disk",
            }
            stage_thresholds = [
                (0.05, "data"),
                (0.15, "features"),
                (0.25, "classifier"),
                (0.50, "regressor"),
                (0.75, "clusterer"),
                (0.90, "saving"),
            ]
            with Progress(
                SpinnerColumn(),
                TextColumn("[progress.description]{task.description}"),
                BarColumn(),
                TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
                TimeElapsedColumn(),
                console=console,
            ) as progress:
                tasks = {}
                for key, label in stages.items():
                    tasks[key] = progress.add_task(f"[dim]{label}", total=100)
                current_stage: Optional[str] = None

                def on_progress(pct, msg):
                    nonlocal current_stage
                    for threshold, stage_key in stage_thresholds:
                        if pct >= threshold and stage_key != current_stage:
                            if current_stage is not None and current_stage in tasks:
                                progress.update(
                                    tasks[current_stage],
                                    completed=100,
                                    description=f"[green]✓ {stages[current_stage]}",
                                )
                            current_stage = stage_key
                            progress.update(
                                tasks[stage_key],
                                completed=50,
                                description=f"[cyan]⧗ {stages[stage_key]}",
                            )
                    overall_pct = int(pct * 100)
                    if current_stage:
                        progress.update(
                            tasks[current_stage], completed=min(95, overall_pct)
                        )

                metrics = pipeline.train_all(
                    data_path,
                    progress_callback=on_progress,
                    n_clusters=n_clusters if train_clu else 5,
                    max_features=max_features,
                )
                # Note: cls_c, reg_alpha, test_split are available for future pipeline enhancements
                for key in tasks:
                    progress.update(
                        tasks[key],
                        completed=100,
                        description=f"[green]✓ {stages[key]}",
                    )
            console.print("\n[bold green]✓ Training Complete![/bold green]")
            mt = Table(title="Training Metrics", box=box.ROUNDED)
            mt.add_column("Model", style="cyan")
            mt.add_column("Metrics", style="yellow")
            for name, vals in metrics.items():
                mt.add_row(name, str(vals))
            console.print(mt)
            tlog.info("TRAIN  complete")
        except Exception as e:
            console.print(f"[red]✗ Training failed: {e}[/red]")
            console.print(
                "[dim]Tip: ensure training CSV exists at specified path[/dim]"
            )
            tlog.info("ERROR  training: %s", e)

    # ── embeddings ──────────────────────────────────────────────
    def embed_flow(self):
        console.print("\n[bold yellow]Generate Embeddings (ChromaDB)[/bold yellow]")
        console.print(
            "[dim]Indexes corpus into ChromaDB using sentence-transformers.\n"
            "Enables semantic search during hypothetical generation.\n"
            "Optional — generation falls back to keyword matching without embeddings.[/dim]\n"
        )
        corpus_path = _path("Corpus JSON to index", default=self._corpus_path)
        if corpus_path is None:
            return
        if not Path(corpus_path).exists():
            console.print(f"[red]✗ File not found: {corpus_path}[/red]")
            return
        if not _confirm(f"Index {corpus_path} into ChromaDB?", default=True):
            return
        try:
            with open(corpus_path) as f:
                data = json.load(f)
            entries = data if isinstance(data, list) else data.get("entries", [])
            if not entries:
                console.print("[red]✗ No entries in corpus[/red]")
                return
            from ..services.vector_service import VectorService

            hypos = []
            for i, e in enumerate(entries):
                text = e.get("text", "")
                topics = e.get("topic", e.get("topics", []))
                if isinstance(topics, str):
                    topics = [t.strip() for t in topics.split(",")]
                hypos.append(
                    {
                        "id": f"entry_{i}",
                        "text": text,
                        "topics": topics,
                        "metadata": e.get("metadata", {}),
                    }
                )
            vs = VectorService()
            with Progress(
                SpinnerColumn(),
                TextColumn("[progress.description]{task.description}"),
                BarColumn(),
                TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
                TimeElapsedColumn(),
                console=console,
            ) as progress:
                task = progress.add_task(
                    "[cyan]Generating embeddings", total=len(hypos)
                )
                batch_size = 20
                for start in range(0, len(hypos), batch_size):
                    batch = hypos[start : start + batch_size]
                    _run_async(vs.index_hypotheticals(batch))
                    progress.update(task, advance=len(batch))
            console.print(
                f"[green]✓ Indexed {len(hypos)} entries into ChromaDB[/green]"
            )
            tlog.info("EMBED  %d entries indexed", len(hypos))
        except Exception as e:
            console.print(f"[red]✗ Embedding failed: {e}[/red]")
            console.print(
                "[dim]Tip: check sentence-transformers and chromadb deps[/dim]"
            )
            tlog.info("ERROR  embedding: %s", e)

    # ── corpus ──────────────────────────────────────────────────
    def corpus_flow(self):
        console.print("\n[bold yellow]Browse Corpus[/bold yellow]")
        # Auto-use default corpus if it exists and has entries
        path = self._corpus_path
        if Path(path).exists():
            try:
                with open(path) as f:
                    data = json.load(f)
                entries = data if isinstance(data, list) else data.get("entries", [])
                if entries:
                    console.print(f"[dim]Using corpus: {path} ({len(entries)} entries)[/dim]")
                else:
                    path = _path("Corpus file path", default=self._corpus_path)
            except Exception:
                path = _path("Corpus file path", default=self._corpus_path)
        else:
            path = _path("Corpus file path", default=self._corpus_path)
        if path is None:
            return
        self._load_corpus(path)
        while True:
            c = _select_quit(
                "Corpus Menu",
                choices=[
                    Choice("View entry", value="1"),
                    Choice("Search", value="2"),
                    Choice("Filter by topic", value="3"),
                    Choice("Export CSV", value="4"),
                    Choice("Export JSON", value="5"),
                    Choice("Preprocess raw", value="6"),
                    Choice("Load different file", value="7"),
                ],
            )
            if c is None:
                return
            if c == "1":
                self._view_entry()
            elif c == "2":
                self._search_corpus()
            elif c == "3":
                self._filter_corpus()
            elif c == "4":
                self._export_corpus("csv")
            elif c == "5":
                self._export_corpus("json")
            elif c == "6":
                self._preprocess_raw()
            elif c == "7":
                path = _path(
                    "Corpus file path",
                    default=(self._loaded_path or self._corpus_path),
                )
                if path is not None:
                    self._load_corpus(path)

    def _load_corpus(self, path):
        try:
            p = Path(path)
            if not p.exists():
                console.print(f"[red]✗ File not found: {path}[/red]")
                return
            ext = p.suffix.lower()
            if ext == ".json":
                self._entries = self._parse_json(p)
            elif ext == ".csv":
                self._entries = self._parse_csv(p)
            elif ext == ".txt":
                self._entries = self._parse_txt(p)
            else:
                console.print(f"[red]✗ Unsupported type: {ext}[/red]")
                return
            self._loaded_path = path
            console.print(
                f"[green]✓ Loaded {len(self._entries)} entries from {p.name}[/green]"
            )
            self._display_entries(self._entries)
            tlog.info("CORPUS  loaded %d entries from %s", len(self._entries), path)
        except Exception as e:
            console.print(f"[red]✗ Load error: {e}[/red]")

    def _display_entries(self, entries, page=0, page_size=20):
        while True:
            start = page * page_size
            end = min(start + page_size, len(entries))
            if start >= len(entries):
                return
            table = Table(
                title=f"Entries [{start+1}-{end} of {len(entries)}]", box=box.ROUNDED
            )
            table.add_column("ID", style="cyan", width=5)
            table.add_column("Topics", style="yellow", width=20)
            table.add_column("Quality", style="yellow", width=8)
            table.add_column("Words", style="yellow", width=6)
            table.add_column("Preview", style="dim")
            for i in range(start, end):
                e = entries[i]
                text = e["text"]
                preview = (text[:60] + "...") if len(text) > 60 else text
                preview = preview.replace("\n", " ")
                table.add_row(
                    str(i),
                    e.get("topics", ""),
                    str(e.get("quality", "N/A")),
                    str(len(text.split())),
                    preview,
                )
            console.print(table)
            has_next = end < len(entries)
            has_prev = page > 0
            if not has_next and not has_prev:
                return
            nav = []
            if has_next:
                nav.append(Choice("Next page →", value="next"))
            if has_prev:
                nav.append(Choice("← Previous page", value="prev"))
            nav.append(Choice("Done", value="done"))
            pick = _select_quit(f"Page {page+1} ({end}/{len(entries)})", choices=nav)
            if pick == "next":
                page += 1
            elif pick == "prev":
                page -= 1
            else:
                return

    def _view_entry(self):
        if not self._entries:
            console.print("[red]No corpus loaded[/red]")
            return
        idx = _text(f"Entry ID (0-{len(self._entries)-1})", default="0")
        if idx is None:
            return
        try:
            i = int(idx)
            e = self._entries[i]
            header = f"Topics: {e.get('topics', '')}\n\n" if e.get("topics") else ""
            console.print(
                Panel(
                    header + e["text"],
                    title=f"Entry {i}",
                    box=box.ROUNDED,
                    border_style="cyan",
                )
            )
        except (ValueError, IndexError):
            console.print("[red]✗ Invalid ID[/red]")

    def _search_corpus(self):
        if not self._entries:
            console.print("[red]No corpus loaded[/red]")
            return
        term = _text("Search term")
        if term is None:
            return
        results = [e for e in self._entries if term.lower() in e["text"].lower()]
        console.print(f"[green]Found {len(results)} matches[/green]")
        if results:
            self._display_entries(results)

    def _filter_corpus(self):
        if not self._entries:
            console.print("[red]No corpus loaded[/red]")
            return
        topic = _select_quit("Filter topic", choices=_topic_choices())
        if topic is None:
            return
        results = [
            e for e in self._entries if topic.lower() in e.get("topics", "").lower()
        ]
        console.print(f"[green]Found {len(results)} matches[/green]")
        if results:
            self._display_entries(results)

    def _export_corpus(self, fmt):
        if not self._entries:
            console.print("[red]No corpus data to export[/red]")
            return
        base = Path(self._loaded_path) if self._loaded_path else Path("export")
        path = str(base.with_suffix(f".export.{fmt}"))
        if not _confirm(f"Export to {path}?", default=True):
            return
        try:
            if fmt == "csv":
                with open(path, "w", newline="") as f:
                    w = csv.writer(f)
                    w.writerow(["id", "text", "topics", "quality"])
                    for i, e in enumerate(self._entries):
                        w.writerow([i, e["text"], e["topics"], e.get("quality", "")])
            else:
                with open(path, "w") as f:
                    json.dump(self._entries, f, indent=2)
            console.print(f"[green]✓ Exported to {path}[/green]")
            tlog.info("EXPORT  %s → %s", fmt, path)
        except Exception as e:
            console.print(f"[red]✗ Export failed: {e}[/red]")

    def _preprocess_raw(self):
        raw_dir = _path("Raw corpus directory", default=self._raw_dir)
        if raw_dir is None:
            return
        output_path = _path("Output corpus JSON path", default=self._corpus_path)
        if output_path is None:
            return
        if not _confirm(f"Preprocess {raw_dir} → {output_path}?", default=True):
            return
        try:
            from ..services.corpus_preprocessor import build_corpus

            with console.status(
                "[bold green]Preprocessing raw corpus...", spinner="dots"
            ):
                count = build_corpus(
                    raw_dir=Path(raw_dir), output_path=Path(output_path)
                )
            self._raw_dir = raw_dir
            self._corpus_path = output_path
            console.print(
                f"[green]✓ Preprocessed {count} entries → {output_path}[/green]"
            )
            self._load_corpus(output_path)
            tlog.info("PREPROCESS  %d entries → %s", count, output_path)
        except Exception as e:
            console.print(f"[red]✗ Preprocess failed: {e}[/red]")
            console.print(
                "[dim]Tip: ensure raw directory exists with subdirectories[/dim]"
            )

    def _parse_json(self, p):
        with open(p) as f:
            data = json.load(f)
        if isinstance(data, list):
            return self._normalize_entries(data)
        if isinstance(data, dict) and "entries" in data:
            return self._normalize_entries(data["entries"])
        return [{"text": json.dumps(data, indent=2), "topics": "", "quality": "N/A"}]

    def _normalize_entries(self, items):
        return [
            {
                "text": e.get("text", str(e)),
                "topics": (
                    ", ".join(e["topic"])
                    if isinstance(e.get("topic"), list)
                    else ", ".join(e.get("topics", []))
                ),
                "quality": e.get("quality_score", "N/A"),
            }
            for e in items
        ]

    def _parse_csv(self, p):
        entries = []
        with open(p) as f:
            for row in csv.DictReader(f):
                entries.append(
                    {
                        "text": row.get("text", row.get("content", "")),
                        "topics": row.get("topic_labels", row.get("topics", "")),
                        "quality": row.get("quality_score", row.get("quality", "N/A")),
                    }
                )
        return entries

    def _parse_txt(self, p):
        return [{"text": p.read_text(), "topics": "", "quality": "N/A"}]

    # ── clean up ─────────────────────────────────────────────────
    def cleanup_flow(self):
        """Let users selectively remove Jikai-generated files and directories."""
        import shutil

        console.print("\n[bold yellow]Clean Up[/bold yellow]")
        console.print("[dim]Select which categories of files to remove.[/dim]\n")

        # Define cleanup categories with paths and descriptions
        categories = {
            "config": {
                "label": "Config — .env, .jikai_state, .jikai_state.bak",
                "paths": [".env", ".jikai_state", ".jikai_state.bak", ".jikai_state.tmp"],
            },
            "models": {
                "label": "ML Models — trained classifiers in models/",
                "paths": ["models/classifier.joblib", "models/vectorizer.joblib",
                          "models/binarizer.joblib", "models/regressor.joblib",
                          "models/clusterer.joblib"],
            },
            "history": {
                "label": "Generation History — data/history.json",
                "paths": ["data/history.json"],
            },
            "embeddings": {
                "label": "Vector Embeddings — chroma_db/",
                "paths": ["chroma_db"],
            },
            "logs": {
                "label": "Logs — logs/",
                "paths": ["logs"],
            },
            "labelled": {
                "label": "Labelled Data — corpus/labelled/",
                "paths": ["corpus/labelled"],
            },
            "db": {
                "label": "Database — data/jikai.db",
                "paths": ["data/jikai.db"],
            },
        }

        choices = [
            Choice(cat["label"], value=key)
            for key, cat in categories.items()
        ]

        selected = _checkbox("Select categories to clean up", choices=choices)
        if not selected:
            console.print("[dim]Nothing selected.[/dim]")
            return

        # Show what will be deleted
        st = Table(box=box.ROUNDED, title="Files to Remove", title_style="bold red")
        st.add_column("Category", style="bold")
        st.add_column("Path", style="cyan")
        st.add_column("Status", style="yellow")
        for key in selected:
            cat = categories[key]
            first = True
            for p in cat["paths"]:
                path = Path(p)
                if path.exists():
                    if path.is_dir():
                        # Count contents
                        count = sum(1 for _ in path.rglob("*") if _.is_file())
                        status = f"exists ({count} files)"
                    else:
                        status = "exists"
                else:
                    status = "[dim]not found[/dim]"
                label = cat["label"].split(" — ")[0] if first else ""
                st.add_row(label, p, status)
                first = False
        console.print(st)

        if not _confirm("Proceed with deletion? This cannot be undone.", default=False):
            console.print("[dim]Cancelled.[/dim]")
            return

        removed = 0
        skipped = 0
        for key in selected:
            cat = categories[key]
            for p in cat["paths"]:
                path = Path(p)
                if path.exists():
                    try:
                        if path.is_dir():
                            # Preserve .gitkeep files
                            gitkeep = path / ".gitkeep"
                            has_gitkeep = gitkeep.exists()
                            shutil.rmtree(path)
                            if has_gitkeep:
                                path.mkdir(parents=True, exist_ok=True)
                                gitkeep.touch()
                            removed += 1
                        else:
                            path.unlink()
                            removed += 1
                    except OSError as e:
                        console.print(f"[red]✗ Failed to remove {p}: {e}[/red]")
                        skipped += 1
                else:
                    skipped += 1

        console.print(f"\n[green]✓ Cleaned up {removed} item(s)[/green]", end="")
        if skipped:
            console.print(f" [dim]({skipped} already absent)[/dim]")
        else:
            console.print()
        tlog.info("CLEANUP  removed=%d skipped=%d categories=%s", removed, skipped, selected)

    # ── settings ────────────────────────────────────────────────
    def settings_flow(self):
        console.print("\n[bold yellow]Settings[/bold yellow]")

        def _mask(v):
            """Masked value for Rich table display."""
            if not v:
                return "[dim]not set[/dim]"
            return v[:4] + "****" + v[-4:] if len(v) > 8 else "****"

        def _mask_plain(v):
            """Masked value for questionary prompts (no Rich markup)."""
            if not v:
                return "not set"
            return v[:4] + "****" + v[-4:] if len(v) > 8 else "****"

        def _render_settings_table(env):
            """Show all current settings grouped by category."""
            st = Table(box=box.ROUNDED, title="Current Settings", title_style="bold cyan")
            st.add_column("Category", style="bold", no_wrap=True)
            st.add_column("Setting", style="cyan")
            st.add_column("Value", style="yellow")
            # API Keys
            st.add_row("API Keys", "Anthropic", _mask(env.get("ANTHROPIC_API_KEY", "")))
            st.add_row("", "OpenAI", _mask(env.get("OPENAI_API_KEY", "")))
            st.add_row("", "Google", _mask(env.get("GOOGLE_API_KEY", "")))
            st.add_row("", "", "")
            # Hosts
            st.add_row("Hosts", "Ollama", env.get("OLLAMA_HOST", "http://localhost:11434"))
            st.add_row("", "Local LLM", env.get("LOCAL_LLM_HOST", "http://localhost:8080"))
            st.add_row("", "", "")
            # Generation Defaults
            st.add_row("Defaults", "Temperature", env.get("DEFAULT_TEMPERATURE", "0.7"))
            st.add_row("", "Max Tokens", env.get("DEFAULT_MAX_TOKENS", "2048"))
            st.add_row("", "", "")
            # Paths & Logging
            st.add_row("Paths", "Corpus", env.get("CORPUS_PATH", self._corpus_path))
            st.add_row("", "Database", env.get("DATABASE_PATH", "data/jikai.db"))
            st.add_row("", "Log Level", env.get("LOG_LEVEL", "INFO"))
            console.print(st)

        def _save_env(env):
            """Write current env dict to .env, preserving all keys."""
            managed_keys = [
                "ANTHROPIC_API_KEY", "OPENAI_API_KEY", "GOOGLE_API_KEY",
                "OLLAMA_HOST", "LOCAL_LLM_HOST",
                "DEFAULT_TEMPERATURE", "DEFAULT_MAX_TOKENS",
                "CORPUS_PATH", "DATABASE_PATH", "LOG_LEVEL",
            ]
            lines = [f"{k}={env[k]}" for k in managed_keys if k in env]
            # Preserve any extra keys the user may have added manually
            for k, v in env.items():
                if k not in managed_keys:
                    lines.append(f"{k}={v}")
            fd = os.open(".env", os.O_WRONLY | os.O_CREAT | os.O_TRUNC, 0o600)
            with os.fdopen(fd, "w") as f:
                f.write("\n".join(lines) + "\n")
            console.print("[green]✓ Settings saved to .env[/green]")
            tlog.info("SETTINGS  saved")

        env = self._load_env()
        _render_settings_table(env)

        while True:
            console.print()
            c = _select_quit(
                "What would you like to edit?",
                choices=[
                    Choice("API Keys — Anthropic, OpenAI, Google", value="keys"),
                    Choice("Hosts — Ollama, Local LLM endpoints", value="hosts"),
                    Choice("Generation Defaults — temperature, max tokens", value="defaults"),
                    Choice("Paths & Logging — corpus, database, log level", value="paths"),
                ],
            )
            if c is None:
                return

            if c == "keys":
                console.print("\n[bold cyan]API Keys[/bold cyan]")
                console.print("[dim]Press Enter to keep current value. Leave blank to clear.[/dim]\n")
                anthropic_key = _password(
                    f"Anthropic API Key ({_mask_plain(env.get('ANTHROPIC_API_KEY', ''))})",
                    default=env.get("ANTHROPIC_API_KEY", ""),
                )
                if anthropic_key is None:
                    continue
                openai_key = _password(
                    f"OpenAI API Key ({_mask_plain(env.get('OPENAI_API_KEY', ''))})",
                    default=env.get("OPENAI_API_KEY", ""),
                )
                if openai_key is None:
                    continue
                google_key = _password(
                    f"Google API Key ({_mask_plain(env.get('GOOGLE_API_KEY', ''))})",
                    default=env.get("GOOGLE_API_KEY", ""),
                )
                if google_key is None:
                    continue
                env["ANTHROPIC_API_KEY"] = anthropic_key
                env["OPENAI_API_KEY"] = openai_key
                env["GOOGLE_API_KEY"] = google_key
                _save_env(env)
                _render_settings_table(env)

            elif c == "hosts":
                console.print("\n[bold cyan]Host Configuration[/bold cyan]")
                console.print("[dim]Press Enter to keep current value.[/dim]\n")
                ollama_host = _text(
                    "Ollama Host",
                    default=env.get("OLLAMA_HOST", "http://localhost:11434"),
                )
                if ollama_host is None:
                    continue
                local_host = _text(
                    "Local LLM Host",
                    default=env.get("LOCAL_LLM_HOST", "http://localhost:8080"),
                )
                if local_host is None:
                    continue
                env["OLLAMA_HOST"] = ollama_host
                env["LOCAL_LLM_HOST"] = local_host
                _save_env(env)
                _render_settings_table(env)

            elif c == "defaults":
                console.print("\n[bold cyan]Generation Defaults[/bold cyan]")
                console.print("[dim]These set the defaults for Quick Generate mode.[/dim]\n")
                temperature = _validated_text(
                    "Default Temperature (0.0–2.0)",
                    default=env.get("DEFAULT_TEMPERATURE", "0.7"),
                    validate=_validate_number(0.0, 2.0, is_float=True),
                )
                if temperature is None:
                    continue
                max_tokens = _validated_text(
                    "Default Max Tokens (1–100000)",
                    default=env.get("DEFAULT_MAX_TOKENS", "2048"),
                    validate=_validate_number(1, 100000),
                )
                if max_tokens is None:
                    continue
                env["DEFAULT_TEMPERATURE"] = temperature
                env["DEFAULT_MAX_TOKENS"] = max_tokens
                _save_env(env)
                _render_settings_table(env)

            elif c == "paths":
                console.print("\n[bold cyan]Paths & Logging[/bold cyan]")
                console.print("[dim]Press Enter to keep current value.[/dim]\n")
                corpus_path = _path(
                    "Corpus Path",
                    default=env.get("CORPUS_PATH", self._corpus_path),
                )
                if corpus_path is None:
                    continue
                db_path = _path(
                    "Database Path",
                    default=env.get("DATABASE_PATH", "data/jikai.db"),
                )
                if db_path is None:
                    continue
                log_level = _select_quit(
                    "Log Level",
                    choices=[
                        Choice("DEBUG — verbose output for troubleshooting", value="DEBUG"),
                        Choice("INFO — standard operational messages", value="INFO"),
                        Choice("WARNING — potential issues only", value="WARNING"),
                        Choice("ERROR — errors only", value="ERROR"),
                    ],
                )
                if log_level is None:
                    continue
                env["CORPUS_PATH"] = corpus_path
                env["DATABASE_PATH"] = db_path
                env["LOG_LEVEL"] = log_level
                _save_env(env)
                _render_settings_table(env)

    def _load_env(self):
        vals: Dict[str, str] = {}
        try:
            with open(".env") as f:
                for line in f:
                    line = line.strip()
                    if "=" in line and not line.startswith("#"):
                        k, v = line.split("=", 1)
                        vals[k.strip()] = v.strip()
        except FileNotFoundError:
            pass
        return vals

    # ── providers ───────────────────────────────────────────────
    def providers_flow(self):
        console.print("\n[bold yellow]LLM Providers[/bold yellow]")
        while True:
            c = _select_quit(
                "Providers",
                choices=[
                    Choice("Check Health", value="1"),
                    Choice("Set Default Provider", value="2"),
                ],
            )
            if c is None:
                return
            if c == "1":
                self._check_health()
            elif c == "2":
                self._set_default_provider()

    def _check_health(self):
        try:
            from ..services.llm_service import llm_service

            async def _health():
                h = await llm_service.health_check()
                m = await llm_service.list_models()
                return h, m

            with console.status(
                "[bold green]Checking provider health...", spinner="dots"
            ):
                health, models = _run_async(_health())
            ht = Table(title="Provider Health", box=box.ROUNDED)
            ht.add_column("Provider", style="cyan")
            ht.add_column("Status", style="yellow")
            ht.add_column("Models", style="dim")
            for name, status in health.items():
                healthy = (
                    status.get("healthy", False) if isinstance(status, dict) else status
                )
                icon = "[green]✓[/green]" if healthy else "[red]✗[/red]"
                ml = models.get(name, [])
                ht.add_row(name, icon, ", ".join(ml[:5]) if ml else "-")
            console.print(ht)
            tlog.info("HEALTH  %s", {k: str(v) for k, v in health.items()})
        except Exception as e:
            console.print(f"[red]✗ Health check failed: {e}[/red]")

    def _ping_provider(self, provider: str) -> bool:
        """Ping a provider to check if it's reachable. Returns True if healthy."""
        try:
            from ..services.llm_service import llm_service

            async def _check():
                h = await llm_service.health_check()
                return h.get(provider, {})

            status = _run_async(_check())
            if isinstance(status, dict):
                return status.get("healthy", False)
            return bool(status)
        except Exception:
            return False

    def _set_default_provider(self):
        provider = _select_quit("Default provider", choices=PROVIDER_CHOICES)
        if provider is None:
            return
        try:
            from ..services.llm_service import llm_service

            llm_service.select_provider(provider)
            console.print(f"[green]✓ Default provider: {provider}[/green]")
            tlog.info("PROVIDER  default=%s", provider)
        except Exception as e:
            console.print(f"[red]✗ Failed: {e}[/red]")

    # ── service startup ──────────────────────────────────────────
    def _ping_ollama(self, host: str = "http://localhost:11434", timeout: float = 1.5) -> bool:
        """Return True if Ollama is reachable at host."""
        import urllib.request
        import urllib.error
        try:
            if not host.startswith(("http://", "https://")):
                return False
            req = urllib.request.Request(f"{host}/api/tags", method="GET")
            with urllib.request.urlopen(req, timeout=timeout):  # nosec B310
                return True
        except Exception:
            return False

    def _ollama_host_from_env(self) -> str:
        """Read OLLAMA_HOST from .env or fall back to localhost."""
        host = "http://localhost:11434"
        env_file = Path(".env")
        if env_file.exists():
            try:
                for line in env_file.read_text().splitlines():
                    line = line.strip()
                    if line.startswith("OLLAMA_HOST="):
                        val = line.split("=", 1)[1].strip()
                        if val:
                            host = val
            except Exception:
                pass
        return host

    def _configured_provider(self) -> str:
        """Best-effort read of the configured LLM provider from state or .env."""
        state = self._load_state()
        provider = state.get("last_config", {}).get("provider", "")
        if provider:
            return provider
        env_file = Path(".env")
        if env_file.exists():
            try:
                for line in env_file.read_text().splitlines():
                    line = line.strip()
                    if line.startswith("DEFAULT_PROVIDER="):
                        val = line.split("=", 1)[1].strip()
                        if val:
                            return val
            except Exception:
                pass
        return "ollama"  # default

    def _start_ollama(self, host: str) -> bool:
        """
        Ensure Ollama is running.  If not, try to launch it in the background.
        Returns True if Ollama is reachable after the attempt.
        """
        import shutil
        import subprocess
        import time

        if self._ping_ollama(host):
            return True

        if shutil.which("ollama") is None:
            console.print(
                "[yellow]⚠ ollama binary not found. Install from https://ollama.com[/yellow]"
            )
            return False

        console.print("[dim]Starting Ollama server in the background...[/dim]")
        try:
            subprocess.Popen(
                ["ollama", "serve"],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                start_new_session=True,
            )
        except Exception as e:
            console.print(f"[red]✗ Could not start Ollama: {e}[/red]")
            return False

        # Wait up to 8 s for Ollama to become reachable
        deadline = time.monotonic() + 8.0
        with console.status("[cyan]Waiting for Ollama to start…", spinner="dots"):
            while time.monotonic() < deadline:
                time.sleep(0.5)
                if self._ping_ollama(host):
                    console.print("[green]✓ Ollama started successfully.[/green]")
                    return True

        console.print("[yellow]⚠ Ollama did not respond in time — generation may fail.[/yellow]")
        return False

    def _start_services(self):
        """
        Ensure all required background services are running.
        Asks for permission to install missing dependencies for each service.
        """
        console.print("\n[bold cyan]Initializing Services...[/bold cyan]")

        services_to_check = {
            "ocr": "Import & Preprocess Corpus (OCR, PDF, DOCX)",
            "train": "ML Model Training (sklearn, pandas)",
            "embed": "Semantic Search indexing (chromadb, torch, embeddings)",
            "gen": "Core Generation (httpx)",
            "export": "Document Export (python-docx)",
        }

        missing_services = [key for key in services_to_check if not self._check_service_deps(key)]

        if missing_services:
            console.print("[yellow]⚠ Some services are missing dependencies.[/yellow]")
            choices = [Choice(services_to_check[key], value=key) for key in missing_services]
            selected = _checkbox("Select services to install dependencies for:", choices=choices)

            if selected:
                for key in selected:
                    self._install_service_deps(key)

            # Warn about unselected
            unselected = set(missing_services) - set(selected or [])
            if unselected:
                console.print("[yellow]⚠ The following functions will be disabled or limited:[/yellow]")
                for key in unselected:
                    console.print(f"  • {services_to_check[key]}")

        provider = self._configured_provider()
        if provider == "ollama":
            host = self._ollama_host_from_env()
            if self._ping_ollama(host):
                console.print(f"[green]✓ Ollama is running at {host}[/green]")
            else:
                self._start_ollama(host)
        elif provider in ["google", "anthropic"]:
            if not self._check_service_deps(provider):
                console.print(f"[yellow]⚠ {provider.title()} provider dependencies are missing.[/yellow]")
                if _confirm(f"Install dependencies for {provider}?", default=True):
                    self._install_service_deps(provider)

    # ── first-run wizard ─────────────────────────────────────────
    def _needs_first_run(self) -> bool:
        """Detect empty state: no .env, no corpus.json, no models."""
        return not (Path(".env").exists() and self._corpus_ready())

    def _check_single_dep(self, pkg_name: str) -> bool:
        """Return True if the package is installed and importable."""
        import importlib.util
        import_name = IMPORT_MAP.get(pkg_name, pkg_name)
        try:
            return importlib.util.find_spec(import_name) is not None
        except (ImportError, ValueError):
            return False

    def _check_service_deps(self, service_key: str) -> bool:
        """Return True if all dependencies for a service are installed."""
        deps = SERVICE_DEPS.get(service_key, [])
        return all(self._check_single_dep(d) for d in deps)

    def _install_service_deps(self, service_key: str):
        """Install dependencies for a specific service."""
        import subprocess
        import sys
        deps = SERVICE_DEPS.get(service_key, [])
        if not deps:
            return True

        console.print(f"[cyan]Installing dependencies for {service_key}...[/cyan]")
        failed = []
        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            BarColumn(),
            TextColumn("{task.completed}/{task.total}"),
            TimeElapsedColumn(),
            console=console,
        ) as progress:
            task = progress.add_task(f"[cyan]Installing {service_key} deps", total=len(deps))
            for pkg in deps:
                progress.update(task, description=f"[cyan]Installing {pkg}")
                # Use --no-cache-dir to ensure fresh install if needed, but --quiet to keep it clean
                result = subprocess.run(
                    [sys.executable, "-m", "pip", "install", pkg, "--quiet"],
                    capture_output=True,
                    text=True,
                )
                if result.returncode != 0:
                    failed.append(pkg)
                progress.advance(task)

        if failed:
            console.print(f"[red]✗ Failed to install: {', '.join(failed)}[/red]")
            return False
        console.print(f"[green]✓ Dependencies for {service_key} installed.[/green]")
        # Re-apply patch if needed after installing chromadb
        if "chromadb" in deps and sys.version_info >= (3, 14):
            self._apply_chromadb_py314_patch()
        return True

    def _deps_installed(self) -> bool:
        """Check if core dependencies are installed."""
        return self._check_service_deps("gen")

    def _install_deps(self):
        """Install core dependencies."""
        self._install_service_deps("gen")

    def _apply_chromadb_py314_patch(self):
        """
        Patch pydantic v1's ModelMetaclass for Python 3.14+ compatibility.

        Python 3.14 stores class-body annotations via __annotate_func__ (PEP 649)
        instead of populating __annotations__ during class execution.  Pydantic v1's
        ModelMetaclass reads namespace['__annotations__'] which is always empty on
        3.14, causing ConfigError for every annotated field.  We inject the evaluated
        annotations before the metaclass processes the namespace.
        """

        try:
            import pydantic.v1.main as _pyd_main

            _orig = _pyd_main.ModelMetaclass.__new__

            def _fixed(mcs, name, bases, namespace, **kwargs):
                if "__annotations__" not in namespace and "__annotate_func__" in namespace:
                    annotate_func = namespace["__annotate_func__"]
                    try:
                        import annotationlib
                        annotations = annotate_func(annotationlib.Format.VALUE)
                    except Exception:
                        try:
                            annotations = annotate_func(1)  # Format.VALUE == 1
                        except Exception:
                            annotations = {}
                    namespace["__annotations__"] = annotations
                return _orig(mcs, name, bases, namespace, **kwargs)

            _pyd_main.ModelMetaclass.__new__ = _fixed
            console.print(
                "[green]✓ Applied pydantic v1 / Python 3.14 compatibility patch.[/green]"
            )
        except Exception as e:
            console.print(
                f"[yellow]⚠ Could not apply pydantic v1 patch: {e}[/yellow]"
            )

    def first_run_wizard(self):
        """Walk user through initial setup."""
        # create .env template if missing
        if not Path(".env").exists():
            template = (
                "# Jikai environment configuration\n"
                "# ANTHROPIC_API_KEY=\n"
                "# OPENAI_API_KEY=\n"
                "# GOOGLE_API_KEY=\n"
                "# OLLAMA_HOST=http://localhost:11434\n"
                "# LOCAL_LLM_HOST=http://localhost:8080\n"
                "# DEFAULT_TEMPERATURE=0.7\n"
                "# DEFAULT_MAX_TOKENS=2048\n"
                "# CORPUS_PATH=corpus/corpus.json\n"
                "# DATABASE_PATH=data/jikai.db\n"
                "# LOG_LEVEL=INFO\n"
            )
            fd = os.open(".env", os.O_WRONLY | os.O_CREAT | os.O_TRUNC, 0o600)
            with os.fdopen(fd, "w") as f:
                f.write(template)
            console.print("[dim]Created .env template — fill in your API keys.[/dim]")

        console.print(
            Panel(
                "[bold cyan]Welcome to Jikai![/bold cyan]\n"
                "[dim]Let's set up your environment. Steps already done will be skipped.[/dim]",
                box=box.DOUBLE,
                border_style="cyan",
            )
        )

        deps_done = self._deps_installed()
        steps = [
            ("Install Dependencies", deps_done, "deps"),
            ("Configure API Keys", Path(".env").exists(), "settings"),
            ("Preprocess Corpus", self._corpus_ready(), "ocr"),
            ("Label Entries (optional)", self._labelled_ready(), "label"),
            ("Generate First Hypothetical", False, "gen"),
        ]
        checklist = Table(title="Setup Checklist", box=box.ROUNDED)
        checklist.add_column("Step", style="cyan")
        checklist.add_column("Status", style="yellow")
        for name, done, _ in steps:
            icon = "[green]✓ Done[/green]" if done else "[dim]Pending[/dim]"
            checklist.add_row(name, icon)
        console.print(checklist)

        for name, done, action in steps:
            if done:
                console.print(f"[green]✓ {name} — already done, skipping[/green]")
                continue
            if not _confirm(f"Proceed with: {name}?", default=True):
                continue
            if action == "deps":
                self._install_deps()
            elif action == "settings":
                # auto-detect Ollama before settings
                try:
                    import urllib.request

                    req = urllib.request.Request(
                        "http://localhost:11434/api/tags", method="GET"
                    )
                    with urllib.request.urlopen(req, timeout=2):  # nosec B310
                        console.print(
                            "[green]✓ Ollama detected at localhost:11434 — available as provider[/green]"
                        )
                except Exception:
                    pass
                self.settings_flow()
                # post-settings provider health check
                try:
                    self._check_health()
                except Exception:
                    pass
            elif action == "ocr":
                self.ocr_flow()
            elif action == "label":
                self.label_flow()
            elif action == "gen":
                if self._corpus_ready():
                    self.generate_flow()
                else:
                    console.print(
                        "[yellow]Corpus not ready yet. Skipping generation.[/yellow]"
                    )
        console.print("[bold green]Setup complete! Entering main menu.[/bold green]\n")

    # ── entry point ─────────────────────────────────────────────
    def run(self):
        self.display_banner()
        self._start_services()
        if self._needs_first_run():
            self.first_run_wizard()
        self.main_menu()


if __name__ == "__main__":
    JikaiTUI().run()
