"""Generation flow."""

from pathlib import Path
from typing import Dict, List

from questionary import Choice
from rich import box
from rich.live import Live
from rich.panel import Panel
from rich.progress import (
    BarColumn,
    Progress,
    SpinnerColumn,
    TextColumn,
    TimeElapsedColumn,
)
from rich.table import Table
from rich.text import Text

from src.tui.console import console, tlog
from src.tui.constants import HISTORY_PATH, TOPICS
from src.tui.inputs import (
    COMPLEXITY_CHOICES,
    METHOD_CHOICES,
    PARTIES_CHOICES,
    PROVIDER_CHOICES,
    _checkbox,
    _confirm,
    _model_choices,
    _path,
    _select_quit,
    _text,
    _topic_choices,
    _validate_number,
    _validated_text,
)
from src.tui.types import GenerationConfig
from src.tui.utils import run_async


class GenerationFlowMixin:
    """Mixin for generation flows."""

    def generate_flow(self):
        while True:
            console.print("\n[bold yellow]Generate Hypothetical[/bold yellow]")
            mode = _select_quit(
                "Mode",
                choices=[
                    Choice("Quick generate — topic only, use defaults", value="quick"),
                    Choice("Custom generate — full config", value="custom"),
                ],
            )
            if mode is None:
                return

            state = self._load_state()
            defaults = state.get("last_config", {})

            if mode == "quick":
                # Show defaults first before committing
                provider = defaults.get("provider", "ollama")
                model_name = defaults.get("model", None)
                temperature = defaults.get("temperature", 0.7)
                complexity = defaults.get("complexity", "3")
                parties = defaults.get("parties", "2")
                method = defaults.get("method", "pure_llm")
                red_herrings = False

                dt = Table(title="Quick Generate Defaults", box=box.SIMPLE)
                dt.add_column("Setting", style="cyan")
                dt.add_column("Value", style="yellow")
                dt.add_row("Provider", provider)
                dt.add_row("Model", model_name or "default")
                dt.add_row("Temperature", str(temperature))
                dt.add_row("Complexity", complexity)
                dt.add_row("Parties", parties)
                dt.add_row("Method", method)
                console.print(dt)

                if not _confirm("Use these settings?", default=True):
                    console.print("[dim]Switching to custom mode...[/dim]")
                    mode = "custom"
                else:
                    topic = _select_quit("Topic", choices=_topic_choices())
                    if topic is None:
                        return

            if mode == "custom":
                topic = _select_quit("Topic", choices=_topic_choices())
                if topic is None:
                    return
                provider = _select_quit("Provider", choices=PROVIDER_CHOICES)
                if provider is None:
                    return
                # Auto-check provider health
                with console.status(f"[dim]Checking {provider}...", spinner="dots"):
                    healthy = self._ping_provider(provider)
                if not healthy:
                    console.print(
                        f"[yellow]⚠ {provider} unreachable — check connection or API key[/yellow]"
                    )
                    if not _confirm(f"Continue with {provider} anyway?", default=False):
                        return
                model_name = _select_quit("Model", choices=_model_choices(provider))
                if model_name is None:
                    return
                if model_name == "__custom__":
                    model_name = _text("Custom model name", default="")
                    if model_name is None:
                        return
                temperature = _select_quit(
                    "Temperature",
                    choices=[
                        Choice(f"{v/10:.1f}", value=str(v / 10)) for v in range(1, 16)
                    ],
                )
                if temperature is None:
                    return
                temperature = float(temperature)
                complexity = _select_quit("Complexity", choices=COMPLEXITY_CHOICES)
                if complexity is None:
                    return
                parties = _select_quit("Parties", choices=PARTIES_CHOICES)
                if parties is None:
                    return
                method = _select_quit("Method", choices=METHOD_CHOICES)
                if method is None:
                    return
                red_herrings = _confirm("Include red herrings?", default=False)

            cfg = Table(box=box.SIMPLE, title="Generation Config")
            cfg.add_column("Parameter", style="cyan")
            cfg.add_column("Value", style="yellow")
            cfg.add_row("Topic", topic)
            cfg.add_row("Provider", provider)
            cfg.add_row("Model", str(model_name) or "(default)")
            cfg.add_row("Temperature", f"{float(temperature):.1f}")
            cfg.add_row("Complexity", str(complexity))
            cfg.add_row("Parties", str(parties))
            cfg.add_row("Method", method)
            cfg.add_row("Red Herrings", "Yes" if red_herrings else "No")
            console.print(cfg)
            if not _confirm("Proceed with generation?", default=True):
                return
            tlog.info(
                "GENERATE  topic=%s provider=%s model=%s temp=%.1f "
                "complexity=%s parties=%s method=%s",
                topic,
                provider,
                model_name,
                temperature,
                complexity,
                parties,
                method,
            )
            self._do_generate(
                GenerationConfig(
                    topic=topic,
                    provider=provider,
                    model=model_name or None,
                    complexity=int(complexity),
                    parties=int(parties),
                    method=method,
                    temperature=float(temperature),
                    red_herrings=red_herrings,
                )
            )
            # Persist last-used config for quick-generate
            state = self._load_state()
            state["last_config"] = {
                "provider": provider,
                "model": model_name,
                "temperature": float(temperature),
                "complexity": str(complexity),
                "parties": str(parties),
                "method": method,
            }
            self._save_state(state)
            if not _confirm("Generate another?", default=False):
                return

    def _do_generate(self, cfg: "GenerationConfig"):
        topic, provider, model = cfg.topic, cfg.provider, cfg.model
        complexity, parties, method = cfg.complexity, cfg.parties, cfg.method
        temperature, red_herrings = cfg.temperature, cfg.red_herrings
        max_retries = 3
        partial_result = ""
        try:
            from src.services.llm_service import LLMRequest, llm_service
            from src.services.validation_service import validation_service

            red_herring_instr = ""
            if red_herrings:
                red_herring_instr = (
                    " Include 1-2 legally irrelevant but plausible facts as red herrings."
                    " The red herrings should not dominate the scenario."
                )
            request = LLMRequest(
                prompt=(
                    f"Generate a Singapore tort law hypothetical about "
                    f"{topic} with {parties} parties at complexity "
                    f"{complexity}/5.{red_herring_instr}"
                ),
                temperature=temperature,
                stream=True,
            )

            async def _stream():
                nonlocal partial_result
                chunks: List[str] = []
                with Live(
                    Text("Generating..."), console=console, refresh_per_second=4
                ) as live:
                    async for chunk in llm_service.stream_generate(
                        request, provider=provider, model=model
                    ):
                        chunks.append(chunk)
                        partial_result = "".join(chunks)
                        live.update(Text(partial_result))
                return partial_result

            try:
                console.print("[dim]Attempt 1/{0}[/dim]".format(max_retries))
                result = run_async(_stream())
                if result:
                    vr = validation_service.validate_hypothetical(
                        text=result,
                        required_topics=[topic],
                        expected_parties=parties,
                    )
                    score = vr.get("overall_score", 0.0)
                    passed = vr.get("passed", False)
                    console.print(
                        Panel(
                            result,
                            title="Generated Hypothetical",
                            box=box.ROUNDED,
                            border_style="green",
                        )
                    )
                    self._show_validation(vr)
                    if passed or score >= 7.0:
                        tlog.info(
                            "GENERATE  stream success, len=%d score=%.1f",
                            len(result),
                            score,
                        )
                        self._save_to_history(
                            {
                                "config": {
                                    "topic": topic,
                                    "provider": provider,
                                    "model": model,
                                    "complexity": complexity,
                                    "parties": parties,
                                    "method": method,
                                },
                                "hypothetical": result,
                                "validation_score": score,
                            }
                        )
                        self._offer_model_answer(
                            result, topic, provider, model, temperature
                        )
                        self._offer_export_now()
                        return
                    retry_reason = self._build_retry_reason(vr, score)
                    console.print(f"[yellow]{retry_reason}[/yellow]")
            except Exception as stream_err:
                tlog.info("GENERATE  stream failed: %s, trying full", stream_err)
                console.print(
                    "[yellow]Stream unavailable "
                    f"({stream_err}), using full generation...[/yellow]"
                )

            from src.services.hypothetical_service import (
                GenerationRequest,
                hypothetical_service,
            )

            feedback = ""
            for attempt in range(1, max_retries + 1):
                console.print(
                    f"[dim]Attempt {attempt}/{max_retries} (full generation)[/dim]"
                )

                async def _full(extra_feedback=feedback):
                    prefs = {"temperature": temperature, "red_herrings": red_herrings}
                    if extra_feedback:
                        prefs["feedback"] = extra_feedback
                    req = GenerationRequest(
                        topics=[topic],
                        number_parties=max(2, min(5, parties)),
                        complexity_level=str(complexity),
                        method=method,
                        provider=provider,
                        model=model,
                        user_preferences=prefs,
                    )
                    return await hypothetical_service.generate_hypothetical(req)

                with console.status(
                    f"[bold green]Generating hypothetical (attempt {attempt}/{max_retries})...",
                    spinner="dots",
                ):
                    response = run_async(_full(feedback))

                partial_result = response.hypothetical
                console.print(
                    Panel(
                        response.hypothetical,
                        title=f"Generated Hypothetical (attempt {attempt})",
                        box=box.ROUNDED,
                        border_style="green",
                    )
                )
                if response.analysis:
                    console.print(
                        Panel(
                            response.analysis,
                            title="Legal Analysis",
                            box=box.ROUNDED,
                            border_style="cyan",
                        )
                    )
                vr = response.validation_results
                self._show_validation(vr)
                score = vr.get("quality_score", 0.0)
                passed = vr.get("passed", False)
                if passed or score >= 7.0:
                    tlog.info(
                        "GENERATE  full success attempt=%d score=%.1f",
                        attempt,
                        score,
                    )
                    self._save_to_history(
                        {
                            "config": {
                                "topic": topic,
                                "provider": provider,
                                "model": model,
                                "complexity": complexity,
                                "parties": parties,
                                "method": method,
                            },
                            "hypothetical": response.hypothetical,
                            "analysis": response.analysis,
                            "validation_score": score,
                        }
                    )
                    self._offer_model_answer(
                        response.hypothetical, topic, provider, model, temperature
                    )
                    self._offer_export_now()
                    return

                if attempt < max_retries:
                    # Build feedback for next attempt
                    checks = vr.get(
                        "checks", vr.get("adherence_check", {}).get("checks", {})
                    )
                    issues = []
                    if isinstance(checks, dict):
                        pc = checks.get("party_count", {})
                        if not pc.get("passed", True):
                            found = pc.get("found", "?")
                            issues.append(
                                f"Include exactly {parties} distinct named parties (found {found})."
                            )
                        ti = checks.get("topic_inclusion", {})
                        if not ti.get("passed", True):
                            missing = ti.get("topics_missing", [])
                            covered = ti.get("topics_found", [])
                            issues.append(
                                f"Must address these topics: {', '.join(missing)}."
                                + (
                                    f" (covered: {', '.join(covered)})"
                                    if covered
                                    else ""
                                )
                            )
                        wc = checks.get("word_count", {})
                        if not wc.get("passed", True):
                            actual = wc.get("count", wc.get("word_count", "?"))
                            issues.append(
                                f"Aim for 800-1500 words (currently {actual})."
                            )
                        if not checks.get("singapore_context", {}).get("passed", True):
                            issues.append("Set the scenario explicitly in Singapore.")
                    feedback = "Previous attempt scored {:.1f}/10. Fix these issues: {}".format(
                        score,
                        " ".join(issues) if issues else "Improve overall quality.",
                    )
                    retry_reason = self._build_retry_reason(vr, score)
                    console.print(f"[yellow]{retry_reason}[/yellow]")

            console.print(
                "[yellow]Max retries reached. Showing best result above.[/yellow]"
            )
            tlog.info("GENERATE  max retries reached, score=%.1f", score)
        except ImportError as e:
            console.print(f"[red]✗ Missing dependency: {e}[/red]")
            console.print(
                "[yellow]Run: pip install chromadb sentence-transformers[/yellow]"
            )
            tlog.info("ERROR  import: %s", e)
        except KeyboardInterrupt:
            word_count = len(partial_result.split()) if partial_result else 0
            if partial_result and word_count > 20:
                console.print(
                    f"\n[yellow]Generation cancelled. Partial result ({word_count} words):[/yellow]"
                )
                console.print(
                    Panel(
                        partial_result,
                        title="Partial Hypothetical",
                        box=box.ROUNDED,
                        border_style="yellow",
                    )
                )
            else:
                console.print("\n[yellow]Generation cancelled.[/yellow]")
            tlog.info("GENERATE  cancelled by user")
        except Exception as e:
            self._handle_generation_error(e)
            tlog.info("ERROR  generation: %s", e)

    def _handle_generation_error(self, e: Exception):
        """Pattern-match common errors and provide specific recovery messages."""
        err_str = str(e).lower()

        # API key errors
        if "api key" in err_str or "apikey" in err_str or "authentication" in err_str:
            if "openai" in err_str:
                console.print(
                    "[red]✗ No API key set for OpenAI. Go to More → Settings to add one.[/red]"
                )
            elif "anthropic" in err_str:
                console.print(
                    "[red]✗ No API key set for Anthropic. Go to More → Settings to add one.[/red]"
                )
            elif "google" in err_str:
                console.print(
                    "[red]✗ No API key set for Google. Go to More → Settings to add one.[/red]"
                )
            else:
                console.print(
                    "[red]✗ API key missing or invalid. Go to More → Settings to check your keys.[/red]"
                )
        # Model not found
        elif "model" in err_str and (
            "not found" in err_str
            or "does not exist" in err_str
            or "not available" in err_str
        ):
            console.print(
                "[red]✗ Model not available. Run More → Providers → Check Health to see available models.[/red]"
            )
        # Connection/network errors — attempt to auto-restart Ollama
        elif (
            "timeout" in err_str
            or "timed out" in err_str
            or "connection" in err_str
            or "refused" in err_str
            or "unreachable" in err_str
            or "all connection attempts failed" in err_str
        ):
            console.print("[red]✗ Cannot connect to provider.[/red]")
            provider_name = self._configured_provider()
            if provider_name == "ollama":
                host = self._ollama_host_from_env()
                started = self._start_ollama(host)
                if started:
                    console.print(
                        "[dim]Ollama is now running — press Generate again to retry.[/dim]"
                    )
                else:
                    console.print(
                        "[dim]Could not start Ollama automatically. "
                        "Run 'ollama serve' in a terminal, then try again.[/dim]"
                    )
            else:
                console.print(
                    "[dim]Check that your provider service is running and accessible.[/dim]"
                )
        # Rate limiting
        elif "rate limit" in err_str or "too many requests" in err_str:
            console.print(
                "[red]✗ Rate limited by provider. Wait a moment and try again.[/red]"
            )
        # Generic fallback
        else:
            console.print(f"[red]✗ Generation failed: {e}[/red]")
            console.print("[dim]Tip: check provider health via More → Providers[/dim]")

    def _offer_model_answer(
        self, hypothetical_text, topic, provider, model, temperature
    ):
        """Offer to generate a model answer (IRAC analysis) for the hypothetical."""
        if not _confirm("Generate model answer?", default=False):
            return
        try:
            from src.services.llm_service import LLMRequest, llm_service

            prompt = (
                "You are an expert Singapore tort law tutor. Given the following hypothetical, "
                "produce a model answer consisting of:\n"
                "1. ISSUE-SPOTTING CHECKLIST: List every legal issue present.\n"
                "2. STRUCTURED IRAC ANALYSIS: For each issue, provide:\n"
                "   - Issue: State the legal question\n"
                "   - Rule: State the relevant legal principle(s) and SG case law\n"
                "   - Application: Apply the rule to the facts\n"
                "   - Conclusion: State the likely outcome\n\n"
                f"HYPOTHETICAL:\n{hypothetical_text}"
            )
            req = LLMRequest(
                prompt=prompt,
                temperature=max(0.3, temperature - 0.2),
                max_tokens=3000,
            )

            async def _gen():
                return await llm_service.generate(req, provider=provider, model=model)

            with console.status(
                "[bold green]Generating model answer...", spinner="dots"
            ):
                resp = run_async(_gen())
            console.print(
                Panel(
                    resp.content,
                    title="Model Answer",
                    box=box.ROUNDED,
                    border_style="magenta",
                )
            )
            tlog.info("MODEL_ANSWER  generated, len=%d", len(resp.content))
        except Exception as e:
            console.print(f"[red]✗ Model answer generation failed: {e}[/red]")

    def _show_validation(self, vr):
        """Display validation results table with detailed check results."""
        if not vr:
            return
        vt = Table(title="Validation Results", box=box.ROUNDED)
        vt.add_column("Check", style="cyan")
        vt.add_column("Result", style="yellow")

        # Extract detailed checks
        checks = vr.get("checks", vr.get("adherence_check", {}).get("checks", {}))
        if isinstance(checks, dict):
            # Party Count
            pc = checks.get("party_count", {})
            if pc:
                passed = pc.get("passed", True)
                found = pc.get("found", "?")
                expected = pc.get("expected", "?")
                icon = "[green]✓[/green]" if passed else "[red]✗[/red]"
                vt.add_row("Party Count", f"{icon} Found {found} (expected {expected})")

            # Topic Coverage
            ti = checks.get("topic_inclusion", {})
            if ti:
                passed = ti.get("passed", True)
                icon = "[green]✓[/green]" if passed else "[red]✗[/red]"
                if passed:
                    vt.add_row("Topic Coverage", f"{icon} Covered")
                else:
                    missing = ti.get("missing", ti.get("topics_missing", []))
                    if missing:
                        vt.add_row(
                            "Topic Coverage", f"{icon} Missing: {', '.join(missing)}"
                        )
                    else:
                        vt.add_row("Topic Coverage", f"{icon} Incomplete")

            # Word Count
            wc = checks.get("word_count", {})
            if wc:
                passed = wc.get("passed", True)
                count = wc.get("count", wc.get("word_count", "?"))
                icon = "[green]✓[/green]" if passed else "[red]✗[/red]"
                vt.add_row("Word Count", f"{icon} {count} words (target: 800-1500)")

            # Singapore Context
            sc = checks.get("singapore_context", {})
            if sc:
                passed = sc.get("passed", True)
                icon = "[green]✓[/green]" if passed else "[red]✗[/red]"
                vt.add_row(
                    "Singapore Context", f"{icon} {'Present' if passed else 'Missing'}"
                )

        # Quality Score
        score = vr.get("quality_score", vr.get("overall_score", "N/A"))
        if isinstance(score, (int, float)):
            vt.add_row("Quality Score", f"{score:.1f} / 10")
        else:
            vt.add_row("Quality Score", str(score))

        # Overall
        passed = vr.get("passed", False)
        vt.add_row(
            "Overall",
            "[green]✓ Passed[/green]" if passed else "[red]✗ Failed[/red]",
        )
        console.print(vt)

    def _build_retry_reason(self, vr, score):
        """Build a detailed retry reason message from validation results."""
        checks = vr.get("checks", vr.get("adherence_check", {}).get("checks", {}))
        failed_checks = []
        if isinstance(checks, dict):
            pc = checks.get("party_count", {})
            if not pc.get("passed", True):
                found = pc.get("found", "?")
                expected = pc.get("expected", "?")
                failed_checks.append(f"party count ({found}/{expected})")
            ti = checks.get("topic_inclusion", {})
            if not ti.get("passed", True):
                missing = ti.get("missing", [])
                if missing:
                    failed_checks.append(
                        f"topic coverage (missing: {', '.join(missing)})"
                    )
                else:
                    failed_checks.append("topic coverage")
            wc = checks.get("word_count", {})
            if not wc.get("passed", True):
                actual = wc.get("count", wc.get("word_count", "?"))
                failed_checks.append(f"word count ({actual} words)")
            sc = checks.get("singapore_context", {})
            if not sc.get("passed", True):
                failed_checks.append("Singapore context")
        if failed_checks:
            return f"Score {score:.1f}/10 — minimum is 7.0, retrying (failed: {', '.join(failed_checks)})..."
        return f"Score {score:.1f}/10 — minimum is 7.0, retrying with corrective feedback..."

    def _offer_export_now(self):
        """Offer to export the last generated hypothetical immediately."""
        if not _confirm("Export now?", default=False):
            return
        self.export_flow()

    def batch_generate_flow(self):
        """Generate multiple hypotheticals in batch."""
        console.print("\n[bold yellow]Batch Generate Hypotheticals[/bold yellow]")
        count = _validated_text(
            "Number to generate (2-10)", default="3", validate=_validate_number(2, 10)
        )
        if count is None:
            return
        count = int(count)
        strategy = _select_quit(
            "Topic spread",
            choices=[
                Choice("Random — pick topics randomly", value="random"),
                Choice("Specified — choose topics", value="specified"),
                Choice("Balanced — cycle through all topics", value="balanced"),
            ],
        )
        if strategy is None:
            return
        topics_list: List[str] = []
        if strategy == "specified":
            selected = _checkbox("Select topics", choices=_topic_choices())
            if not selected:
                return
            topics_list = selected
        elif strategy == "balanced":
            topics_list = list(TOPICS)
        elif strategy == "random":
            import random

            topics_list = random.sample(TOPICS, min(count, len(TOPICS)))

        provider = _select_quit("Provider", choices=PROVIDER_CHOICES)
        if provider is None:
            return
        # Auto-check provider health
        with console.status(f"[dim]Checking {provider}...", spinner="dots"):
            healthy = self._ping_provider(provider)
        if not healthy:
            console.print(
                f"[yellow]⚠ {provider} unreachable — check connection or API key[/yellow]"
            )
            if not _confirm(f"Continue with {provider} anyway?", default=False):
                return
        model_name = _select_quit("Model", choices=_model_choices(provider))
        if model_name is None:
            return
        if model_name == "__custom__":
            model_name = _text("Custom model name", default="")
            if model_name is None:
                return
        complexity = _select_quit("Complexity", choices=COMPLEXITY_CHOICES)
        if complexity is None:
            return
        parties = _select_quit("Parties", choices=PARTIES_CHOICES)
        if parties is None:
            return
        method = _select_quit("Method", choices=METHOD_CHOICES)
        if method is None:
            return
        if not _confirm(f"Generate {count} hypotheticals?", default=True):
            return

        from src.services.hypothetical_service import (
            GenerationRequest,
            hypothetical_service,
        )

        results = []
        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            BarColumn(),
            TextColumn("{task.completed}/{task.total}"),
            TimeElapsedColumn(),
            console=console,
        ) as progress:
            task = progress.add_task("[cyan]Generating batch", total=count)
            for i in range(count):
                # Pick topic based on strategy
                if strategy == "balanced":
                    topic = topics_list[i % len(topics_list)]
                elif strategy == "specified":
                    topic = topics_list[i % len(topics_list)]
                else:
                    import random

                    topic = random.choice(topics_list)
                progress.update(
                    task,
                    description=f"[cyan]Generating {i+1}/{count}: {topic}",
                )
                try:

                    async def _gen(t=topic):
                        req = GenerationRequest(
                            topics=[t],
                            number_parties=max(2, min(5, int(parties))),
                            complexity_level=str(complexity),
                            method=method,
                            provider=provider,
                            model=model_name or None,
                        )
                        return await hypothetical_service.generate_hypothetical(req)

                    resp = run_async(_gen(topic))
                    score = resp.validation_results.get("quality_score", 0.0)
                    results.append(
                        {
                            "topic": topic,
                            "score": score,
                            "words": len(resp.hypothetical.split()),
                        }
                    )
                except Exception as e:
                    results.append(
                        {"topic": topic, "score": 0.0, "words": 0, "error": str(e)}
                    )
                progress.advance(task)

        # Summary table
        st = Table(title=f"Batch Results ({count} hypotheticals)", box=box.ROUNDED)
        st.add_column("#", style="cyan", width=3)
        st.add_column("Topic", style="yellow")
        st.add_column("Score", style="green")
        st.add_column("Words", style="dim")
        st.add_column("Status", style="cyan")
        for i, r in enumerate(results, 1):
            status = (
                "[green]OK[/green]" if r.get("score", 0) >= 7.0 else "[red]LOW[/red]"
            )
            if "error" in r:
                status = f"[red]ERR: {r['error'][:30]}[/red]"
            st.add_row(str(i), r["topic"], f"{r['score']:.1f}", str(r["words"]), status)
        console.print(st)
        tlog.info("BATCH  generated %d hypotheticals", count)

    def export_flow(self):
        """Export a generated hypothetical + analysis to DOCX (and optionally PDF)."""
        console.print("\n[bold yellow]Export Hypothetical to DOCX/PDF[/bold yellow]")
        console.print(
            "[dim]Export a generated hypothetical to a document.\n"
            "Requires python-docx (pip install python-docx).\n"
            "PDF conversion requires docx2pdf (optional).[/dim]\n"
        )
        history = self._load_history()
        hypo_text = ""
        analysis_text = ""
        model_answer = ""

        # Offer options: export from history or paste manually
        source_choices = []
        if history:
            source_choices.append(Choice("Export last generated", value="last"))
            source_choices.append(Choice("Select from history", value="history"))
        source_choices.append(Choice("Paste text manually", value="paste"))
        source_choices.append(Choice("Back", value="back"))

        source = _select_quit("Export source", choices=source_choices)
        if source is None or source == "back":
            return

        if source == "last" and history:
            rec = history[-1]
            hypo_text = rec.get("hypothetical", "")
            analysis_text = rec.get("analysis", "")
            model_answer = rec.get("model_answer", "")
            topic = rec.get("config", {}).get("topic", "hypothetical")
            console.print(f"[green]✓ Loaded last generated ({topic})[/green]")
        elif source == "history" and history:
            # Show history selection
            ht = Table(
                title=f"Select from History ({len(history)} records)", box=box.ROUNDED
            )
            ht.add_column("#", style="cyan", width=4)
            ht.add_column("Time", style="dim", width=16)
            ht.add_column("Topic", style="yellow", width=18)
            ht.add_column("Preview", style="dim")
            display_records = history[-20:]
            for i, r in enumerate(display_records, 1):
                ts = r.get("timestamp", "")[:16]
                cfg = r.get("config", {})
                topic = cfg.get("topic", str(cfg.get("topics", "")))
                text = r.get("hypothetical", "")[:60]
                ht.add_row(str(i), ts, topic, text.replace("\n", " "))
            console.print(ht)
            idx = _text(f"Select record # (1-{len(display_records)})", default="1")
            if not idx:
                return
            try:
                rec = display_records[int(idx) - 1]
                hypo_text = rec.get("hypothetical", "")
                analysis_text = rec.get("analysis", "")
                model_answer = rec.get("model_answer", "")
                console.print(f"[green]✓ Loaded record #{idx}[/green]")
            except (ValueError, IndexError):
                console.print("[red]Invalid selection[/red]")
                return
        elif source == "paste":
            hypo_text = _text(
                "Paste hypothetical text (or path to .txt file)", default=""
            )
            if not hypo_text:
                return
            # Check if it's a file path
            if Path(hypo_text).exists() and Path(hypo_text).is_file():
                hypo_text = Path(hypo_text).read_text(encoding="utf-8")
            analysis_text = _text(
                "Paste analysis text (optional, Enter to skip)", default=""
            )
            model_answer = _text(
                "Paste model answer (optional, Enter to skip)", default=""
            )

        if not hypo_text:
            console.print("[red]✗ No hypothetical text to export[/red]")
            return

        out_path = _path("Output DOCX path", default="exports/hypothetical.docx")
        if not out_path:
            return
        also_pdf = _confirm("Also export PDF?", default=False)
        try:
            import docx
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt

            doc = docx.Document()
            style = doc.styles["Normal"]
            font = style.font
            font.name = "Times New Roman"
            font.size = Pt(12)

            doc.add_heading("Legal Hypothetical", level=1)
            for i, para in enumerate(hypo_text.split("\n\n")):
                p = doc.add_paragraph(para.strip())
                if i == 0:
                    p.paragraph_format.space_before = Pt(12)

            if analysis_text:
                doc.add_heading("Legal Analysis", level=1)
                for para in analysis_text.split("\n\n"):
                    doc.add_paragraph(para.strip())

            if model_answer:
                doc.add_heading("Model Answer", level=1)
                for para in model_answer.split("\n\n"):
                    doc.add_paragraph(para.strip())

            # Footer with metadata
            from datetime import datetime

            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = (
                f"Generated by Jikai | {datetime.now().strftime('%Y-%m-%d %H:%M')}"
            )
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            Path(out_path).parent.mkdir(parents=True, exist_ok=True)
            doc.save(out_path)
            console.print(f"[green]✓ DOCX saved → {out_path}[/green]")

            if also_pdf:
                try:
                    from docx2pdf import convert

                    pdf_path = str(Path(out_path).with_suffix(".pdf"))
                    convert(out_path, pdf_path)
                    console.print(f"[green]✓ PDF saved → {pdf_path}[/green]")
                except ImportError:
                    console.print(
                        "[yellow]docx2pdf not installed. Install with: pip install docx2pdf[/yellow]"
                    )
            tlog.info("EXPORT  %s", out_path)
        except ImportError:
            console.print(
                "[red]✗ python-docx not installed. Run: pip install python-docx[/red]"
            )
        except Exception as e:
            console.print(f"[red]✗ Export failed: {e}[/red]")
            tlog.info("ERROR  export: %s", e)

    def _offer_variation(self, original_record: Dict):
        """Offer to create a variation of a generated or historical hypothetical."""
        if not _confirm("Create a variation?", default=False):
            return
        cfg = original_record.get("config", {})
        param = _select_quit(
            "Which parameter to vary?",
            choices=[
                Choice("Swap topic", value="topic"),
                Choice("Change party count", value="parties"),
                Choice("Change complexity", value="complexity"),
                Choice("Change difficulty/method", value="method"),
            ],
        )
        if param is None:
            return
        new_cfg = dict(cfg)
        if param == "topic":
            new_topic = _select_quit("New topic", choices=_topic_choices())
            if new_topic is None:
                return
            new_cfg["topic"] = new_topic
        elif param == "parties":
            new_parties = _select_quit("New party count", choices=PARTIES_CHOICES)
            if new_parties is None:
                return
            new_cfg["parties"] = int(new_parties)
        elif param == "complexity":
            new_c = _select_quit("New complexity", choices=COMPLEXITY_CHOICES)
            if new_c is None:
                return
            new_cfg["complexity"] = int(new_c)
        elif param == "method":
            new_m = _select_quit("New method", choices=METHOD_CHOICES)
            if new_m is None:
                return
            new_cfg["method"] = new_m

        # Generate variation using full generation with reference to original
        try:
            from src.services.hypothetical_service import (
                GenerationRequest,
                hypothetical_service,
            )

            original_text = original_record.get("hypothetical", "")
            prefs = {
                "feedback": (
                    f"Create a variation of this hypothetical by changing the {param}. "
                    f"Original scenario (for reference, do NOT copy):\n{original_text[:500]}..."
                )
            }

            async def _var():
                req = GenerationRequest(
                    topics=[new_cfg.get("topic", cfg.get("topic", "negligence"))],
                    number_parties=max(2, min(5, new_cfg.get("parties", 3))),
                    complexity_level=str(new_cfg.get("complexity", "3")),
                    method=new_cfg.get("method", "pure_llm"),
                    provider=new_cfg.get("provider"),
                    model=new_cfg.get("model"),
                    user_preferences=prefs,
                )
                return await hypothetical_service.generate_hypothetical(req)

            with console.status("[bold green]Generating variation...", spinner="dots"):
                resp = run_async(_var())
            console.print(
                Panel(
                    resp.hypothetical,
                    title=f"Variation (changed: {param})",
                    box=box.ROUNDED,
                    border_style="yellow",
                )
            )
            score = resp.validation_results.get("quality_score", 0.0)
            self._save_to_history(
                {
                    "config": new_cfg,
                    "hypothetical": resp.hypothetical,
                    "analysis": resp.analysis,
                    "validation_score": score,
                    "variation_of": original_record.get("timestamp", ""),
                }
            )
            console.print(f"[green]✓ Variation saved (score: {score:.1f})[/green]")
            tlog.info("VARIATION  success score=%.1f param=%s", score, param)
        except Exception as e:
            console.print(f"[red]✗ Variation failed: {e}[/red]")
            tlog.error("VARIATION  failed: %s", e)
