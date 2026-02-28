"""
FastAPI application for Jikai legal hypothetical generation service.
Provides REST API endpoints for generating, validating, and managing legal hypotheticals.
"""

import os
import time
import uuid
from collections import defaultdict
from datetime import datetime
from typing import Any, Dict, List, Optional

import structlog
from fastapi import (
    BackgroundTasks,
    Depends,
    FastAPI,
    HTTPException,
    Query,
    Request,
    status,
)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.trustedhost import TrustedHostMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field
from starlette.middleware.base import BaseHTTPMiddleware
from structlog.contextvars import bind_contextvars, clear_contextvars

from ..config import settings
from ..domain import canonicalize_topic, is_tort_topic
from ..services import (
    CorpusQuery,
    CorpusService,
    GenerationReport,
    GenerationRequest,
    GenerationResponse,
    HypotheticalEntry,
    HypotheticalService,
    LLMService,
    corpus_service,
    database_service,
    hypothetical_service,
    llm_service,
)
from ..services.topic_guard import canonicalize_and_validate_topics
from .web import web_router

# Configure structured logging
structlog.configure(
    processors=[
        structlog.contextvars.merge_contextvars,
        structlog.stdlib.filter_by_level,
        structlog.stdlib.add_logger_name,
        structlog.stdlib.add_log_level,
        structlog.stdlib.PositionalArgumentsFormatter(),
        structlog.processors.TimeStamper(fmt="iso"),
        structlog.processors.StackInfoRenderer(),
        structlog.processors.format_exc_info,
        structlog.processors.UnicodeDecoder(),
        structlog.processors.JSONRenderer(),
    ],
    context_class=dict,
    logger_factory=structlog.stdlib.LoggerFactory(),
    wrapper_class=structlog.stdlib.BoundLogger,
    cache_logger_on_first_use=True,
)

logger = structlog.get_logger(__name__)

# Create FastAPI application
app = FastAPI(
    title="Jikai Legal Hypothetical Generator",
    description="AI-powered service for generating Singapore Tort Law hypotheticals for educational purposes",
    version=settings.app_version,
    docs_url="/docs" if settings.api.debug else None,
    redoc_url="/redoc" if settings.api.debug else None,
)

app.include_router(web_router)

# Add middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=settings.api.cors_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.add_middleware(
    TrustedHostMiddleware,
    allowed_hosts=["*"] if settings.api.debug else ["localhost", "127.0.0.1"],
)


# API Key authentication middleware
JIKAI_API_KEY = os.environ.get("JIKAI_API_KEY", "")
_EXEMPT_PATHS = {"/", "/health", "/docs", "/redoc", "/openapi.json"}
REQUEST_ID_HEADER = "X-Request-ID"


class RequestIDMiddleware(BaseHTTPMiddleware):
    """Attach a request ID to contextvars and response headers for each request."""

    async def dispatch(self, request: Request, call_next):
        request_id = request.headers.get(REQUEST_ID_HEADER, "").strip() or str(
            uuid.uuid4()
        )
        request.state.request_id = request_id
        bind_contextvars(request_id=request_id)
        started = time.perf_counter()
        try:
            response = await call_next(request)
        except Exception:
            duration_ms = round((time.perf_counter() - started) * 1000, 2)
            logger.exception(
                "request_unhandled_exception",
                method=request.method,
                path=request.url.path,
                duration_ms=duration_ms,
            )
            clear_contextvars()
            raise

        duration_ms = round((time.perf_counter() - started) * 1000, 2)
        response.headers[REQUEST_ID_HEADER] = request_id
        logger.info(
            "request_completed",
            method=request.method,
            path=request.url.path,
            status_code=response.status_code,
            duration_ms=duration_ms,
        )
        clear_contextvars()
        return response


class APIKeyMiddleware(BaseHTTPMiddleware):
    """Validate JIKAI_API_KEY header on every request (except exempt paths)."""

    async def dispatch(self, request: Request, call_next):
        if not JIKAI_API_KEY:
            # No key configured — skip auth
            return await call_next(request)
        if request.url.path in _EXEMPT_PATHS:
            return await call_next(request)
        key = request.headers.get("X-API-Key", "")
        if key != JIKAI_API_KEY:
            client_ip = request.client.host if request.client else "unknown"
            logger.warning("auth_failure", path=request.url.path, client_ip=client_ip)
            return JSONResponse(
                status_code=401,
                content={"error": "Invalid or missing API key. Set X-API-Key header."},
            )
        return await call_next(request)


# Token-bucket rate limiter middleware
class RateLimiterMiddleware(BaseHTTPMiddleware):
    """Token-bucket rate limiter per client IP."""

    def __init__(self, app, rate_limit: int = 100, window: int = 60):
        super().__init__(app)
        self.rate_limit = rate_limit
        self.window = window
        self._buckets: Dict[str, list] = defaultdict(list)

    async def dispatch(self, request: Request, call_next):
        client_ip = request.client.host if request.client else "unknown"
        now = time.time()
        # Prune old timestamps
        self._buckets[client_ip] = [
            ts for ts in self._buckets[client_ip] if now - ts < self.window
        ]
        if len(self._buckets[client_ip]) >= self.rate_limit:
            logger.warning("rate_limited", client_ip=client_ip)
            return JSONResponse(
                status_code=429,
                content={
                    "error": f"Rate limit exceeded. Max {self.rate_limit} requests per {self.window}s."
                },
            )
        self._buckets[client_ip].append(now)
        # evict oldest IPs if tracking too many
        if len(self._buckets) > 10000:
            oldest_ip = min(
                self._buckets,
                key=lambda ip: self._buckets[ip][0] if self._buckets[ip] else 0,
            )
            del self._buckets[oldest_ip]
        return await call_next(request)


app.add_middleware(APIKeyMiddleware)
app.add_middleware(
    RateLimiterMiddleware,
    rate_limit=settings.api.rate_limit,
    window=60,
)
app.add_middleware(RequestIDMiddleware)


# Request/Response Models
class HealthResponse(BaseModel):
    """Health check response model."""

    status: str
    timestamp: str
    version: str
    services: Dict[str, Any]


class TopicsResponse(BaseModel):
    """Available topics response model."""

    topics: List[str]
    count: int


class GenerationStatsResponse(BaseModel):
    """Generation statistics response model."""

    total_generations: int
    average_generation_time: float
    success_rate: float
    latency_metrics: Dict[str, Any] = Field(default_factory=dict)
    recent_generations: List[Dict[str, Any]]


# Dependency functions
async def get_hypothetical_service():
    """Dependency to get hypothetical service instance."""
    return hypothetical_service


async def get_llm_service():
    """Dependency to get LLM service instance."""
    return llm_service


async def get_corpus_service():
    """Dependency to get corpus service instance."""
    return corpus_service


# Startup and shutdown events
@app.on_event("startup")
async def startup_event():
    """Application startup event."""
    logger.info(
        "Starting Jikai API",
        version=settings.app_version,
        environment=settings.environment,
    )

    # Initialize services
    try:
        migrated_count = await database_service.migrate_legacy_history_json()
        if migrated_count:
            logger.info("Migrated legacy history.json into SQLite", count=migrated_count)
        # Health check all services
        health_status = await hypothetical_service.health_check()
        logger.info("Services initialized", health_status=health_status)
    except Exception as e:
        logger.error("Failed to initialize services", error=str(e))
        raise


@app.on_event("shutdown")
async def shutdown_event():
    """Application shutdown event."""
    logger.info("Shutting down Jikai API")

    # Close service connections
    try:
        await llm_service.close()
        logger.info("Services closed successfully")
    except Exception as e:
        logger.error("Error during shutdown", error=str(e))


# Health check endpoint
@app.get("/health", response_model=HealthResponse)
async def health_check():
    """Health check endpoint for monitoring."""
    try:
        # Check all services
        hypothetical_health = await hypothetical_service.health_check()
        llm_health = await llm_service.health_check()
        corpus_health = await corpus_service.health_check()

        # Determine overall status
        overall_status = "healthy"
        if hypothetical_health.get("status") != "healthy":
            overall_status = "unhealthy"
        elif any(
            dep.get("status") == "unhealthy"
            for dep in hypothetical_health.get("dependencies", {}).values()
        ):
            overall_status = "degraded"

        return HealthResponse(
            status=overall_status,
            timestamp=datetime.utcnow().isoformat(),
            version=settings.app_version,
            services={
                "hypothetical_service": hypothetical_health,
                "llm_service": llm_health,
                "corpus_service": corpus_health,
            },
        )
    except Exception as e:
        logger.error("Health check failed", error=str(e))
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=f"Health check failed: {e}",
        )


# Main generation endpoint
@app.post("/generate", response_model=GenerationResponse)
async def generate_hypothetical(
    request: GenerationRequest,
    background_tasks: BackgroundTasks,
    http_request: Request,
    service: HypotheticalService = Depends(get_hypothetical_service),
):
    """Generate a legal hypothetical with analysis."""
    try:
        correlation_id = (
            request.correlation_id
            or getattr(http_request.state, "request_id", None)
            or str(uuid.uuid4())
        )
        logger.info(
            "Hypothetical generation requested",
            topics=request.topics,
            parties=request.number_parties,
            correlation_id=correlation_id,
        )

        canonical_topics = canonicalize_and_validate_topics(request.topics)

        # Validate topics
        topic_extraction_started = time.perf_counter()
        available_topics = [
            canonicalize_topic(topic)
            for topic in await corpus_service.extract_all_topics()
        ]
        topic_extraction_time_ms = round(
            (time.perf_counter() - topic_extraction_started) * 1000, 2
        )
        request = request.model_copy(
            update={
                "topics": canonical_topics,
                "correlation_id": correlation_id,
                "topic_extraction_time_ms": topic_extraction_time_ms,
            }
        )
        invalid_topics = [
            topic for topic in request.topics if topic not in available_topics
        ]

        if invalid_topics:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Invalid topics: {invalid_topics}. Available topics: {available_topics[:10]}...",
            )

        # Generate hypothetical
        response = await service.generate_hypothetical(request)

        # Log generation in background
        background_tasks.add_task(log_generation_event, request.dict(), response.dict())

        logger.info(
            "Hypothetical generated successfully",
            generation_time=response.generation_time,
            validation_passed=response.validation_results.get("passed", False),
            correlation_id=correlation_id,
        )

        return response

    except HTTPException:
        raise
    except Exception as e:
        logger.error(
            "Hypothetical generation failed",
            error=str(e),
            correlation_id=correlation_id,
        )
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Generation failed: {e}",
        )


# Batch generation endpoint
class BatchGenerateConfig(BaseModel):
    topic: str
    provider: Optional[str] = None
    model: Optional[str] = None
    complexity: str = "intermediate"
    parties: int = 3
    method: str = "pure_llm"


class BatchGenerateRequest(BaseModel):
    configs: List[BatchGenerateConfig] = Field(..., min_items=1, max_items=10)


class BatchGenerateResult(BaseModel):
    hypothetical: str
    analysis: str
    validation_score: float


class GeneratePreviewResponse(BaseModel):
    topics: List[str]
    provider: Optional[str] = None
    model: Optional[str] = None
    estimated_input_tokens: int
    estimated_output_tokens: int
    estimated_total_tokens: int
    estimated_latency_seconds: float
    estimated_cost_usd: float
    confidence: str = "heuristic"


class GenerateReportRequest(BaseModel):
    issue_types: List[str] = Field(..., min_items=1, max_items=10)
    comment: Optional[str] = Field(default=None, max_length=2000)


class GenerateReportResponse(BaseModel):
    report_id: int
    generation_id: int
    issue_types: List[str]
    comment: Optional[str] = None
    timestamp: str


class RegenerateGenerationResponse(BaseModel):
    source_generation_id: int
    feedback_context: str = ""
    regenerated: GenerationResponse


def _estimate_generation_preview(request: GenerationRequest) -> GeneratePreviewResponse:
    """Estimate generation cost/latency without executing generation."""
    complexity_map = {
        "beginner": 2,
        "basic": 2,
        "intermediate": 3,
        "advanced": 4,
    }
    raw_complexity = str(request.complexity_level).strip().lower()
    if raw_complexity.isdigit():
        complexity_factor = max(1, min(5, int(raw_complexity)))
    else:
        complexity_factor = complexity_map.get(raw_complexity, 3)

    topic_count = max(1, len(request.topics))
    sample_size = max(1, int(request.sample_size))
    party_count = max(2, int(request.number_parties))

    estimated_input_tokens = (
        650
        + (topic_count * 90)
        + (sample_size * 220)
        + (party_count * 55)
        + (complexity_factor * 120)
    )
    estimated_output_tokens = (
        700
        + (complexity_factor * 260)
        + (party_count * 80)
        + (120 if request.method in ("hybrid", "ml_assisted") else 0)
    )
    total_tokens = estimated_input_tokens + estimated_output_tokens

    provider = (request.provider or settings.llm.provider or "ollama").lower()
    provider_latency_factor = {
        "openai": 1.0,
        "anthropic": 1.1,
        "google": 0.9,
        "ollama": 1.8,
        "local": 1.7,
    }.get(provider, 1.2)
    estimated_latency_seconds = round(
        1.2 + (total_tokens / 900.0) * provider_latency_factor, 2
    )

    # Rough blended token pricing by provider (USD per 1k tokens).
    provider_rate_per_1k = {
        "openai": 0.005,
        "anthropic": 0.006,
        "google": 0.003,
        "ollama": 0.0,
        "local": 0.0,
    }.get(provider, 0.004)
    estimated_cost_usd = round((total_tokens / 1000.0) * provider_rate_per_1k, 6)

    return GeneratePreviewResponse(
        topics=request.topics,
        provider=request.provider,
        model=request.model,
        estimated_input_tokens=estimated_input_tokens,
        estimated_output_tokens=estimated_output_tokens,
        estimated_total_tokens=total_tokens,
        estimated_latency_seconds=estimated_latency_seconds,
        estimated_cost_usd=estimated_cost_usd,
    )


@app.post("/generate/preview", response_model=GeneratePreviewResponse)
async def preview_generation(
    request: GenerationRequest,
):
    """Validate request and return latency/cost metadata without generation."""
    canonical_topics = canonicalize_and_validate_topics(request.topics)
    request = request.model_copy(update={"topics": canonical_topics})

    available_topics = {
        canonicalize_topic(topic) for topic in await corpus_service.extract_all_topics()
    }
    invalid_topics = [topic for topic in request.topics if topic not in available_topics]
    if invalid_topics:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=(
                f"Invalid topics: {invalid_topics}. "
                f"Available topics: {list(sorted(available_topics))[:10]}..."
            ),
        )

    return _estimate_generation_preview(request)


@app.post("/generate/batch")
async def batch_generate(
    request: BatchGenerateRequest,
    http_request: Request,
    service: HypotheticalService = Depends(get_hypothetical_service),
):
    """Generate multiple hypotheticals sequentially. Max 10 per request."""
    batch_correlation_id = (
        getattr(http_request.state, "request_id", None) or str(uuid.uuid4())
    )
    # validate topics upfront
    try:
        available_topics = {
            canonicalize_topic(topic)
            for topic in await corpus_service.extract_all_topics()
        }
    except Exception:
        available_topics = set()
    if available_topics:
        invalid = [
            (i, cfg.topic, canonicalize_topic(cfg.topic))
            for i, cfg in enumerate(request.configs)
            if not is_tort_topic(canonicalize_topic(cfg.topic))
            or canonicalize_topic(cfg.topic) not in available_topics
        ]
        if invalid:
            detail = "; ".join(
                f"config[{i}]: unknown topic '{raw}' (canonical '{canonical}')"
                for i, raw, canonical in invalid
            )
            raise HTTPException(status_code=400, detail=detail)
    results = []
    for idx, cfg in enumerate(request.configs[:10], 1):
        try:
            canonical_topic = canonicalize_and_validate_topics([cfg.topic])[0]
            correlation_id = f"{batch_correlation_id}-batch-{idx}"
            gen_req = GenerationRequest(
                topics=[canonical_topic],
                number_parties=cfg.parties,
                complexity_level=cfg.complexity,
                method=cfg.method,
                provider=cfg.provider,
                model=cfg.model,
                correlation_id=correlation_id,
            )
            resp = await service.generate_hypothetical(gen_req)
            results.append(
                {
                    "hypothetical": resp.hypothetical,
                    "analysis": resp.analysis,
                    "validation_score": resp.validation_results.get(
                        "quality_score", 0.0
                    ),
                }
            )
        except Exception as e:
            results.append(
                {
                    "hypothetical": "",
                    "analysis": "",
                    "validation_score": 0.0,
                    "error": str(e),
                }
            )
    return {"results": results, "count": len(results)}


@app.post(
    "/generate/{generation_id}/report",
    response_model=GenerateReportResponse,
)
async def report_generation_failure(
    generation_id: int,
    request: GenerateReportRequest,
):
    """Persist a generation failure report for regeneration workflows."""
    if generation_id <= 0:
        raise HTTPException(status_code=400, detail="generation_id must be positive")

    from src.services.database_service import database_service

    report = GenerationReport(
        generation_id=generation_id,
        issue_types=request.issue_types,
        comment=request.comment,
        is_locked=True,
    )
    try:
        report_id = await database_service.save_generation_report(report)
    except Exception as e:
        if "FOREIGN KEY constraint failed" in str(e):
            raise HTTPException(
                status_code=404, detail=f"Generation ID {generation_id} not found"
            ) from e
        logger.error("Failed to save generation report", error=str(e))
        raise HTTPException(
            status_code=500, detail=f"Failed to save generation report: {e}"
        ) from e

    return GenerateReportResponse(
        report_id=report_id,
        generation_id=generation_id,
        issue_types=request.issue_types,
        comment=request.comment,
        timestamp=datetime.utcnow().isoformat(),
    )


@app.put("/generate/{generation_id}/report/{report_id}")
async def update_generation_report(generation_id: int, report_id: int):
    """Generation reports are immutable once submitted."""
    raise HTTPException(
        status_code=403,
        detail=(
            f"Report {report_id} for generation {generation_id} is immutable. "
            "Submit a new report instead."
        ),
    )


@app.delete("/generate/{generation_id}/report/{report_id}")
async def delete_generation_report(generation_id: int, report_id: int):
    """Generation reports are immutable once submitted."""
    raise HTTPException(
        status_code=403,
        detail=(
            f"Report {report_id} for generation {generation_id} is immutable. "
            "Deletion is not allowed."
        ),
    )


@app.post(
    "/generate/{generation_id}/regenerate",
    response_model=RegenerateGenerationResponse,
)
async def regenerate_generation(
    generation_id: int,
    http_request: Request,
    service: HypotheticalService = Depends(get_hypothetical_service),
):
    """Force regeneration using original request and stored feedback context."""
    if generation_id <= 0:
        raise HTTPException(status_code=400, detail="generation_id must be positive")

    from src.services.database_service import database_service

    row = await database_service.get_generation_by_id(generation_id)
    if not row:
        raise HTTPException(
            status_code=404, detail=f"Generation ID {generation_id} not found"
        )

    request_data = row.get("request", {})
    original_topics = request_data.get("topics", [])
    if not original_topics:
        raise HTTPException(
            status_code=400,
            detail=f"Generation ID {generation_id} has no stored topics to regenerate",
        )

    canonical_topics = canonicalize_and_validate_topics(original_topics)
    feedback_context = await database_service.build_regeneration_feedback_context(
        generation_id
    )
    reports = await database_service.get_generation_reports(generation_id)
    latest_report = reports[-1] if reports else None
    retry_reason = "report_feedback"
    if latest_report and latest_report.issue_types:
        retry_reason = "report_feedback:" + ",".join(latest_report.issue_types[:3])
    raw_retry_attempt = request_data.get("retry_attempt", 0)
    try:
        retry_attempt = max(1, int(raw_retry_attempt) + 1)
    except (TypeError, ValueError):
        retry_attempt = 1

    user_preferences = request_data.get("user_preferences") or {}
    if feedback_context:
        existing_feedback = str(user_preferences.get("feedback", "")).strip()
        user_preferences["feedback"] = (
            f"{existing_feedback} {feedback_context}".strip()
            if existing_feedback
            else feedback_context
        )

    regenerate_request = GenerationRequest(
        topics=canonical_topics,
        law_domain=request_data.get("law_domain", "tort"),
        number_parties=request_data.get("number_parties", 3),
        complexity_level=request_data.get("complexity_level", "intermediate"),
        sample_size=request_data.get("sample_size", 3),
        user_preferences=user_preferences,
        method=request_data.get("method", "pure_llm"),
        provider=request_data.get("provider"),
        model=request_data.get("model"),
        parent_generation_id=generation_id,
        retry_reason=retry_reason,
        retry_attempt=retry_attempt,
        correlation_id=(
            getattr(http_request.state, "request_id", None) or str(uuid.uuid4())
        ),
    )

    regenerated = await service.generate_hypothetical(regenerate_request)
    return RegenerateGenerationResponse(
        source_generation_id=generation_id,
        feedback_context=feedback_context,
        regenerated=regenerated,
    )


# Topics endpoint
@app.get("/topics", response_model=TopicsResponse)
async def get_available_topics(service: CorpusService = Depends(get_corpus_service)):
    """Get all available legal topics from the corpus."""
    try:
        topics = await service.extract_all_topics()

        return TopicsResponse(topics=topics, count=len(topics))

    except Exception as e:
        logger.error("Failed to get topics", error=str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to get topics: {e}",
        )


# Corpus management endpoints
@app.get("/corpus/entries")
async def get_corpus_entries(
    topics: Optional[List[str]] = None,
    limit: int = Query(default=10, ge=1, le=100),
    service: CorpusService = Depends(get_corpus_service),
):
    """Get corpus entries with optional topic filtering."""
    try:
        if topics:
            query = CorpusQuery(topics=topics, sample_size=limit)
            entries = await service.query_relevant_hypotheticals(query)
        else:
            # Get all entries (limited)
            all_entries = await service.load_corpus()
            entries = all_entries[:limit]

        return {"entries": [entry.dict() for entry in entries], "count": len(entries)}

    except Exception as e:
        logger.error("Failed to get corpus entries", error=str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to get corpus entries: {e}",
        )


@app.post("/corpus/entries")
async def add_corpus_entry(
    entry: HypotheticalEntry, service: CorpusService = Depends(get_corpus_service)
):
    """Add a new entry to the corpus."""
    try:
        entry_id = await service.add_hypothetical(entry)

        logger.info("Corpus entry added", id=entry_id, topics=entry.topics)

        return {"id": entry_id, "message": "Entry added successfully"}

    except Exception as e:
        logger.error("Failed to add corpus entry", error=str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to add entry: {e}",
        )


# LLM service endpoints
@app.get("/llm/models")
async def get_available_models(service: LLMService = Depends(get_llm_service)):
    """Get available LLM models."""
    try:
        models = await service.list_models()
        return {"models": models}

    except Exception as e:
        logger.error("Failed to get models", error=str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to get models: {e}",
        )


@app.get("/llm/health")
async def llm_health_check(service: LLMService = Depends(get_llm_service)):
    """Check LLM service health."""
    try:
        health_status = await service.health_check()
        return {"health": health_status}

    except Exception as e:
        logger.error("LLM health check failed", error=str(e))
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=f"LLM health check failed: {e}",
        )


# Statistics endpoint
@app.get("/stats", response_model=GenerationStatsResponse)
async def get_generation_stats(
    service: HypotheticalService = Depends(get_hypothetical_service),
):
    """Get generation statistics from database."""
    try:
        # Get statistics from database (persistent across restarts)
        from src.services.database_service import database_service

        db_stats = await database_service.get_statistics()

        # Get recent generations
        recent = await service.get_generation_history(limit=10)

        return GenerationStatsResponse(
            total_generations=db_stats["total_generations"],
            average_generation_time=db_stats["average_generation_time"],
            success_rate=db_stats["success_rate"],
            latency_metrics=db_stats.get("latency_metrics", {}),
            recent_generations=recent,
        )

    except Exception as e:
        logger.error("Failed to get statistics", error=str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to get statistics: {e}",
        )


# ML endpoints
@app.post("/ml/train")
async def train_ml_models(
    data_path: Optional[str] = None,
    n_clusters: int = 5,
):
    """Trigger ML model training."""
    try:
        from ..ml.pipeline import MLPipeline

        pipeline = MLPipeline()
        import asyncio

        path = data_path or settings.ml.training_data_path
        metrics = await asyncio.to_thread(
            pipeline.train_all, path, n_clusters=n_clusters
        )
        return {"status": "trained", "metrics": metrics}
    except Exception as e:
        logger.error("ML training failed", error=str(e))
        raise HTTPException(status_code=500, detail=f"Training failed: {e}")


@app.get("/ml/status")
async def ml_status():
    """Get ML model status and metrics."""
    try:
        from ..ml.pipeline import MLPipeline

        pipeline = MLPipeline()
        pipeline.load_all()
        return pipeline.get_status()
    except Exception as e:
        return {
            "classifier_trained": False,
            "regressor_trained": False,
            "clusterer_trained": False,
            "error": str(e),
        }


# Export endpoint
@app.get("/export/{generation_id}")
async def export_generation(
    generation_id: int, format: str = "docx", background_tasks: BackgroundTasks = None
):
    """Export a generation by SQLite generation ID as DOCX or PDF file download."""
    if format not in ("docx", "pdf"):
        raise HTTPException(status_code=400, detail="Format must be 'docx' or 'pdf'")
    if generation_id <= 0:
        raise HTTPException(
            status_code=400, detail="generation_id must be a positive integer"
        )

    await database_service.migrate_legacy_history_json()
    row = await database_service.get_generation_by_id(generation_id)
    if not row:
        raise HTTPException(
            status_code=404, detail=f"Generation ID {generation_id} not found"
        )

    response_record = row.get("response", {})
    hypo = response_record.get("hypothetical", "")
    analysis = response_record.get("analysis", "")
    model_answer = response_record.get("model_answer", "")

    try:
        import tempfile

        import docx
        from docx.shared import Pt

        doc = docx.Document()
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        doc.add_heading("Legal Hypothetical", level=1)
        for para in hypo.split("\n\n"):
            doc.add_paragraph(para.strip())
        if analysis:
            doc.add_heading("Legal Analysis", level=1)
            for para in analysis.split("\n\n"):
                doc.add_paragraph(para.strip())
        if model_answer:
            doc.add_heading("Model Answer", level=1)
            for para in model_answer.split("\n\n"):
                doc.add_paragraph(para.strip())

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as docx_tmp:
            docx_path = docx_tmp.name
        doc.save(docx_path)
        temp_paths = [docx_path]
        resp_path = docx_path

        if format == "pdf":
            try:
                from docx2pdf import convert

                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as pdf_tmp:
                    pdf_path = pdf_tmp.name
                convert(docx_path, pdf_path)
                temp_paths.append(pdf_path)
                resp_path = pdf_path
            except ImportError:
                try:
                    os.unlink(docx_path)
                except OSError:
                    pass
                raise HTTPException(
                    status_code=500, detail="docx2pdf not installed for PDF conversion"
                )
            except Exception:
                for path in temp_paths:
                    try:
                        os.unlink(path)
                    except OSError:
                        pass
                raise

        from fastapi.responses import FileResponse

        media = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        if format == "pdf":
            media = "application/pdf"

        def _cleanup_tmp(paths: List[str]):
            for path in paths:
                try:
                    os.unlink(path)
                except OSError:
                    pass

        if background_tasks:
            background_tasks.add_task(_cleanup_tmp, temp_paths)
        return FileResponse(
            resp_path,
            media_type=media,
            filename=f"hypothetical_{generation_id}.{format}",
        )
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export failed: {e}")


@app.get("/providers")
async def list_providers():
    """List all providers and their models."""
    try:
        models = await llm_service.list_models()
        health = await llm_service.health_check()
        return {
            "providers": {
                name: {"models": m, "health": health.get(name, {})}
                for name, m in models.items()
            }
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


class SetDefaultProviderRequest(BaseModel):
    provider: str
    model: Optional[str] = None


@app.put("/providers/default")
async def set_default_provider(req: SetDefaultProviderRequest):
    """Set default LLM provider and model."""
    try:
        llm_service.select_provider(req.provider)
        if req.model:
            llm_service.select_model(req.model)
        return {"default_provider": req.provider, "default_model": req.model}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


# Error handlers
@app.exception_handler(HTTPException)
async def http_exception_handler(request, exc):
    """Handle HTTP exceptions."""
    logger.warning(
        "HTTP exception",
        status_code=exc.status_code,
        detail=exc.detail,
        path=request.url.path,
    )

    return JSONResponse(
        status_code=exc.status_code,
        content={
            "error": exc.detail,
            "status_code": exc.status_code,
            "timestamp": datetime.utcnow().isoformat(),
        },
    )


@app.exception_handler(Exception)
async def general_exception_handler(request, exc):
    """Handle general exceptions."""
    logger.error("Unhandled exception", error=str(exc), path=request.url.path)

    return JSONResponse(
        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
        content={
            "error": "Internal server error",
            "status_code": 500,
            "timestamp": datetime.utcnow().isoformat(),
        },
    )


# Background task functions
async def log_generation_event(
    request_data: Dict[str, Any], response_data: Dict[str, Any]
):
    """Log generation event for analytics."""
    try:
        # This could be extended to log to a database or analytics service
        logger.info(
            "Generation event logged",
            topics=request_data.get("topics"),
            generation_time=response_data.get("generation_time"),
            validation_passed=response_data.get("validation_results", {}).get("passed"),
        )
    except Exception as e:
        logger.error("Failed to log generation event", error=str(e))


# Root endpoint
@app.get("/")
async def root():
    """Root endpoint with API information."""
    return {
        "name": "Jikai Legal Hypothetical Generator",
        "version": settings.app_version,
        "description": "AI-powered service for generating Singapore Tort Law hypotheticals",
        "docs_url": (
            "/docs"
            if settings.api.debug
            else "Documentation not available in production"
        ),
        "health_url": "/health",
    }


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(
        "src.api.main:app",
        host=settings.api.host,
        port=settings.api.port,
        reload=settings.api.debug,
        log_level=settings.logging.level.lower(),
    )
