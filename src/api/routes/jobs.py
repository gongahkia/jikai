"""Long-running job endpoints (preprocess, scrape, train, embed, export, cleanup)."""

import asyncio
import uuid
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, HTTPException
from pydantic import BaseModel, Field

router = APIRouter()

_jobs: Dict[str, Dict[str, Any]] = {}  # in-memory job store


class JobType(str, Enum):
    preprocess = "preprocess"
    scrape = "scrape"
    train = "train"
    embed = "embed"
    export = "export"
    cleanup = "cleanup"
    label = "label"


class PreprocessRequest(BaseModel):
    raw_dir: Optional[str] = None
    output_path: Optional[str] = None
    merge_existing: bool = True
    include_non_tort: bool = False


class ScrapeRequest(BaseModel):
    source: str  # commonlii, judiciary, sicc, gazette
    courts: Optional[List[str]] = None
    years: Optional[List[int]] = None
    max_cases: int = Field(default=50, ge=1, le=500)
    tort_only: bool = True


class TrainRequest(BaseModel):
    data_path: str = "corpus/labelled/sample.csv"
    models: List[str] = Field(default=["classifier", "regressor", "clusterer"])
    n_clusters: int = Field(default=5, ge=2, le=20)


class EmbedRequest(BaseModel):
    corpus_path: str = "corpus/clean/tort/corpus.json"
    batch_size: int = Field(default=20, ge=1, le=100)


class ExportRequest(BaseModel):
    generation_id: Optional[int] = None
    hypothetical: Optional[str] = None
    analysis: Optional[str] = None
    model_answer: Optional[str] = None
    format: str = "docx"  # docx or pdf
    output_path: Optional[str] = None


class CleanupRequest(BaseModel):
    targets: List[
        str
    ]  # config, models, history, embeddings, logs, labelled, database, tui_state


class LabelEntry(BaseModel):
    text: str
    topics: List[str]
    quality_score: float = Field(default=5.0, ge=0, le=10)
    difficulty_level: str = "medium"


class LabelRequest(BaseModel):
    corpus_path: str = "corpus/clean/tort/corpus.json"
    output_path: str = "corpus/labelled/sample.csv"
    entries: List[LabelEntry] = Field(default_factory=list)


def _create_job(job_type: str) -> str:
    job_id = str(uuid.uuid4())[:8]
    _jobs[job_id] = {
        "type": job_type,
        "status": "running",
        "progress": 0,
        "result": None,
        "error": None,
    }
    return job_id


@router.post("/preprocess")
async def preprocess(req: PreprocessRequest):
    job_id = _create_job("preprocess")

    async def run():
        try:
            from ...services.corpus_preprocessor import build_corpus

            count = build_corpus(
                raw_dir=Path(req.raw_dir) if req.raw_dir else None,
                output_path=Path(req.output_path) if req.output_path else None,
                merge_existing=req.merge_existing,
                include_non_tort=req.include_non_tort,
            )
            _jobs[job_id]["status"] = "completed"
            _jobs[job_id]["result"] = {"count": count}
        except Exception as e:
            _jobs[job_id]["status"] = "failed"
            _jobs[job_id]["error"] = str(e)

    asyncio.create_task(run())
    return {"job_id": job_id}


@router.post("/scrape")
async def scrape(req: ScrapeRequest):
    job_id = _create_job("scrape")

    async def run():
        try:
            from ...services.scraper_service import run_scraper

            entries = await run_scraper(
                source=req.source,
                courts=req.courts,
                years=req.years,
                max_cases=req.max_cases,
                tort_only=req.tort_only,
            )
            _jobs[job_id]["status"] = "completed"
            _jobs[job_id]["result"] = {"entries": entries, "count": len(entries)}
        except Exception as e:
            _jobs[job_id]["status"] = "failed"
            _jobs[job_id]["error"] = str(e)

    asyncio.create_task(run())
    return {"job_id": job_id}


@router.post("/train")
async def train(req: TrainRequest):
    job_id = _create_job("train")

    async def run():
        try:
            from ...ml.pipeline import MLPipeline

            pipeline = MLPipeline()
            metrics = pipeline.train_all(
                data_path=req.data_path, n_clusters=req.n_clusters
            )
            _jobs[job_id]["status"] = "completed"
            _jobs[job_id]["result"] = {"metrics": metrics}
        except Exception as e:
            _jobs[job_id]["status"] = "failed"
            _jobs[job_id]["error"] = str(e)

    asyncio.create_task(run())
    return {"job_id": job_id}


@router.post("/embed")
async def embed(req: EmbedRequest):
    job_id = _create_job("embed")

    async def run():
        try:
            import json

            from ...services import vector_service

            with open(req.corpus_path, encoding="utf-8") as f:
                corpus = json.load(f)
            entries = [
                {
                    "id": f"entry_{i}",
                    "text": e.get("text", ""),
                    "topics": e.get("topics", e.get("topic", [])),
                    "metadata": e.get("metadata", {}),
                }
                for i, e in enumerate(corpus)
            ]
            count = await vector_service.index_hypotheticals(entries)
            _jobs[job_id]["status"] = "completed"
            _jobs[job_id]["result"] = {"indexed_count": count}
        except Exception as e:
            _jobs[job_id]["status"] = "failed"
            _jobs[job_id]["error"] = str(e)

    asyncio.create_task(run())
    return {"job_id": job_id}


@router.post("/export")
async def export(req: ExportRequest):
    job_id = _create_job("export")

    async def run():
        try:
            output = req.output_path or f"data/export_{job_id}.{req.format}"
            model_answer = req.model_answer or ""
            if req.generation_id:
                from ...services import database_service

                gen = await database_service.get_generation_by_id(req.generation_id)
                if gen:
                    resp = gen.get("response", {})
                    hypo = resp.get("hypothetical", gen.get("hypothetical", req.hypothetical or ""))
                    analysis = resp.get("analysis", gen.get("analysis", req.analysis or ""))
                    model_answer = model_answer or resp.get("model_answer", "")
                else:
                    hypo = req.hypothetical or ""
                    analysis = req.analysis or ""
            else:
                hypo = req.hypothetical or ""
                analysis = req.analysis or ""
            try:
                from docx import Document

                doc = Document()
                doc.add_heading("Jikai -- Generated Hypothetical", 0)
                if hypo:
                    doc.add_heading("Hypothetical", level=1)
                    doc.add_paragraph(hypo)
                if analysis:
                    doc.add_heading("Legal Analysis", level=1)
                    doc.add_paragraph(analysis)
                if model_answer:
                    doc.add_heading("Model Answer", level=1)
                    doc.add_paragraph(model_answer)
                Path(output).parent.mkdir(parents=True, exist_ok=True)
                doc.save(output)
                _jobs[job_id]["status"] = "completed"
                _jobs[job_id]["result"] = {"output_path": output}
            except ImportError:
                _jobs[job_id]["status"] = "failed"
                _jobs[job_id]["error"] = "python-docx not installed"
        except Exception as e:
            _jobs[job_id]["status"] = "failed"
            _jobs[job_id]["error"] = str(e)

    asyncio.create_task(run())
    return {"job_id": job_id}


class ExportTrainingDataRequest(BaseModel):
    output_path: str = "data/generated/approved_training_data.csv"
    min_score: float = Field(default=7.0, ge=0, le=10)


@router.post("/export-training-data")
async def export_training_data(req: ExportTrainingDataRequest):
    from ...services import database_service

    try:
        count = await database_service.export_approved_training_data(
            req.output_path, min_score=req.min_score
        )
        return {"exported_rows": count, "output_path": req.output_path}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


class AnkiExportRequest(BaseModel):
    generation_ids: Optional[List[int]] = None
    output_path: str = "data/export/anki_cards.tsv"
    include_model_answer: bool = True


@router.post("/export-anki")
async def export_anki(req: AnkiExportRequest):
    from ...services import database_service
    from ...services.export_service import export_to_anki_tsv

    try:
        generations = []
        if req.generation_ids:
            for gid in req.generation_ids:
                gen = await database_service.get_generation_by_id(gid)
                if gen:
                    resp = gen.get("response", {})
                    generations.append({
                        "hypothetical": resp.get("hypothetical", ""),
                        "analysis": resp.get("analysis", ""),
                        "model_answer": resp.get("model_answer", ""),
                        "topics": gen.get("request", {}).get("topics", []),
                    })
        else:
            history = await database_service.get_generation_history(limit=100)
            for gen in history:
                resp = gen.get("response", {})
                generations.append({
                    "hypothetical": resp.get("hypothetical", gen.get("hypothetical", "")),
                    "analysis": resp.get("analysis", gen.get("analysis", "")),
                    "model_answer": resp.get("model_answer", ""),
                    "topics": gen.get("request", {}).get("topics", gen.get("topics", [])),
                })
        count = export_to_anki_tsv(generations, req.output_path, req.include_model_answer)
        return {"exported_cards": count, "output_path": req.output_path}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/cleanup")
async def cleanup(req: CleanupRequest):
    import shutil

    removed = []
    target_map = {
        "config": [".env", ".jikai_state.json", "data/tui_state.json", "data/tui.json"],
        "models": ["models/"],
        "history": ["data/history.json"],
        "embeddings": ["chroma_db/"],
        "logs": ["logs/"],
        "labelled": ["corpus/labelled/"],
        "database": ["data/jikai.db"],
        "tui_state": ["data/tui_state.json", "data/tui.json"],
    }
    for target in req.targets:
        for path_str in target_map.get(target, []):
            p = Path(path_str)
            if p.is_dir():
                shutil.rmtree(p, ignore_errors=True)
                removed.append(path_str)
            elif p.is_file():
                p.unlink(missing_ok=True)
                removed.append(path_str)
    return {"removed": removed}


@router.post("/label")
async def label(req: LabelRequest):
    import csv

    output = Path(req.output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    write_header = not output.exists()
    with open(output, "a", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        if write_header:
            writer.writerow(
                ["text", "topic_labels", "quality_score", "difficulty_level"]
            )
        for entry in req.entries:
            writer.writerow(
                [
                    entry.text,
                    "|".join(entry.topics),
                    entry.quality_score,
                    entry.difficulty_level,
                ]
            )
    return {"labelled_count": len(req.entries), "output_path": str(output)}


@router.get("/{job_id}/status")
async def job_status(job_id: str):
    job = _jobs.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    return job


@router.post("/{job_id}/cancel")
async def cancel_job(job_id: str):
    job = _jobs.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    job["status"] = "cancelled"
    return {"cancelled": True}
